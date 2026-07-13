from __future__ import annotations

from copy import deepcopy
from datetime import date, datetime
import json
import os
import re
import shutil
import signal
import subprocess
import sys
import tempfile
import unicodedata
import time
from pathlib import Path
from urllib.parse import urlparse
from zipfile import ZIP_DEFLATED, ZipFile

import streamlit as st
from lxml import etree

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph

from config.constants import OUTPUT_DIR
from services.json_service import guardar_datos_json
from services.openai_service import generar_json_openai
from utils.nombres_archivo import safe_filename
from utils.textos import limpiar_texto
from utils.validaciones import validar_campos_obligatorios


VERSION_INFORME_TECNICO_FINAL = (
    "VERSION_APROBADA_NORMATIVIDAD_COLOMBIA_INTERNACIONAL"
)
CODIGO_FORMATO_INFORME = "GCDTP-F-023 V01"
NOMBRE_PLANTILLA_INFORME = "GCDTP-F-023_V01_Formato_Informe_Final.docx"

TIPOS_PROYECTO_INFORME = [
    "Software",
    "Diseño industrial",
    "Electrónica y automatización",
    "Inteligencia artificial",
    "Desarrollo de marca e identidad visual",
    "Prototipo agroindustrial",
    "Ficha técnica",
    "Otro",
]

CLASIFICACIONES_INFORMACION = [
    "Pública",
    "Pública Clasificada",
    "Pública Reservada",
]

METODOLOGIAS_DESARROLLO = [
    "Metodologías ágiles",
    "Scrum",
    "Kanban",
    "Modelo en cascada",
    "Modelo espiral",
    "Design Thinking",
    "Doble Diamante",
    "Diseño Centrado en el Usuario (DCU)",
    "Diseño para Manufactura y Ensamble (DFMA)",
    "Lean Startup",
    "Stage-Gate",
    "Desarrollo iterativo de prototipos",
    "Ingeniería de sistemas y modelo V",
    "CRISP-DM para proyectos de datos e inteligencia artificial",
    "DMAIC / Six Sigma",
    "Investigación aplicada y validación experimental",
    "TRIZ para solución inventiva de problemas",
    "Otra",
]

CLAVES_CONTENIDO = [
    "introduccion",
    "planteamiento_problema",
    "objetivo_general",
    "objetivos_especificos",
    "estado_arte_tecnica",
    "metodologia_desarrollo",
    "actividades_corregidas",
    "entregables_corregidos",
    "desarrollo_proyecto",
    "desarrollo_actividades",
    "normatividad_aplicable",
    "normas_aplicables",
    "resultados_obtenidos",
    "analisis_viabilidad",
    "propiedad_transferencia",
    "impacto_proyecto",
    "conclusiones",
    "referencias_bibliograficas",
    "anexos",
]


# =====================================================
# UTILIDADES GENERALES
# =====================================================

def obtener_ruta_plantilla() -> Path:
    raiz_proyecto = Path(__file__).resolve().parents[2]

    candidatos = [
        raiz_proyecto / "resources" / NOMBRE_PLANTILLA_INFORME,
        raiz_proyecto / "recursos" / NOMBRE_PLANTILLA_INFORME,
        Path("resources") / NOMBRE_PLANTILLA_INFORME,
        Path("recursos") / NOMBRE_PLANTILLA_INFORME,
    ]

    for ruta in candidatos:
        if ruta.exists():
            return ruta

    raise FileNotFoundError(
        "No se encontró la plantilla oficial del Informe Final. "
        f"Guárdala como resources/{NOMBRE_PLANTILLA_INFORME}"
    )


def fecha_mes_anio_espanol(fecha: date) -> str:
    meses = {
        1: "Enero",
        2: "Febrero",
        3: "Marzo",
        4: "Abril",
        5: "Mayo",
        6: "Junio",
        7: "Julio",
        8: "Agosto",
        9: "Septiembre",
        10: "Octubre",
        11: "Noviembre",
        12: "Diciembre",
    }
    return f"{meses[fecha.month]} {fecha.year}"


def dividir_lineas(texto: str) -> list[str]:
    texto = str(texto or "").replace(";", "\n")
    elementos = [
        limpiar_texto(linea.strip(" -•\t"))
        for linea in texto.splitlines()
        if limpiar_texto(linea.strip(" -•\t"))
    ]

    if len(elementos) <= 1 and "," in texto:
        elementos = [
            limpiar_texto(item)
            for item in texto.split(",")
            if limpiar_texto(item)
        ]

    return elementos


def normalizar_para_comparacion(texto: str) -> str:
    """Normaliza un texto para detectar coincidencias literales."""
    texto_normalizado = limpiar_texto(texto).casefold()
    texto_normalizado = re.sub(
        r"[^a-záéíóúüñ0-9\s]",
        " ",
        texto_normalizado,
    )
    return re.sub(r"\s+", " ", texto_normalizado).strip()


def contiene_copia_textual(
    texto_generado: str,
    textos_originales: list[str],
    minimo_palabras: int = 6,
) -> bool:
    """
    Detecta si un apartado narrativo reutiliza literalmente fragmentos extensos
    suministrados por el usuario.
    """
    generado = normalizar_para_comparacion(texto_generado)

    if not generado:
        return False

    for original in textos_originales:
        palabras = normalizar_para_comparacion(original).split()

        if not palabras:
            continue

        if len(palabras) < minimo_palabras:
            frase = " ".join(palabras)
            if len(palabras) >= 3 and frase in generado:
                return True
            continue

        for inicio in range(len(palabras) - minimo_palabras + 1):
            fragmento = " ".join(
                palabras[inicio:inicio + minimo_palabras]
            )
            if fragmento in generado:
                return True

    return False


def resultados_modo_prueba_sin_copia(
    datos: dict,
    cantidad_entregables: int,
) -> str:
    """
    Genera un texto de respaldo analítico sin transcribir los entregables.
    """
    cantidad_texto = (
        "un entregable técnico"
        if cantidad_entregables == 1
        else f"{cantidad_entregables} entregables técnicos"
    )

    return (
        f"El cierre del proyecto permitió consolidar {cantidad_texto}, cuya "
        "pertinencia se analiza a partir de su correspondencia con los objetivos "
        "planteados, las actividades desarrolladas y el nivel de madurez "
        f"tecnológica {datos.get('trl_alcanzado', '')}. Los productos obtenidos "
        "representan evidencias concretas del avance alcanzado, debido a que "
        "materializan decisiones de diseño, implementación, integración, "
        "documentación o validación realizadas durante la ejecución. Su valoración "
        "no se limita a confirmar su existencia, sino que considera la utilidad "
        "técnica, la coherencia funcional, la trazabilidad con los requerimientos y "
        "la posibilidad de continuar con procesos posteriores de ajuste, prueba o "
        "escalamiento. La relación entre estos resultados y las actividades permite "
        "identificar cómo cada etapa del desarrollo contribuyó a la construcción de "
        "productos verificables. Asimismo, el análisis reconoce el aporte innovador "
        "del proyecto sin atribuir características, desempeños o validaciones que no "
        "hayan sido documentados. Desde la perspectiva del TRL alcanzado, los "
        "entregables constituyen una base técnica para sustentar el estado actual de "
        "la solución y orientar las acciones necesarias para avanzar hacia un nivel "
        "superior de madurez. La tabla presentada a continuación identifica cada "
        "producto mediante una descripción corregida y mejorada técnicamente, e "
        "incluye un espacio independiente para incorporar los enlaces o soportes de "
        "evidencia correspondientes. De esta manera, el apartado conserva una "
        "redacción analítica, evita repetir literalmente la información ingresada y "
        "mantiene una separación clara entre la explicación de los resultados y la "
        "relación detallada de los entregables."
    )


def dividir_entregables(texto: str) -> list[str]:
    """Convierte el campo de entregables en una lista sin alterar su sentido."""
    contenido = str(texto or "").replace(";", "\n")
    entregables: list[str] = []

    for linea in contenido.splitlines():
        linea = re.sub(
            r"^\s*(?:[-•*]|\d+[.)-]?)\s*",
            "",
            linea,
        )
        linea_limpia = limpiar_texto(linea)

        if linea_limpia:
            entregables.append(linea_limpia)

    return entregables


def dividir_actividades(texto: str) -> list[str]:
    """Convierte un solo campo de actividades en una lista limpia."""
    contenido = str(texto or "").replace(";", "\n")
    actividades: list[str] = []

    for linea in contenido.splitlines():
        linea = re.sub(
            r"^\s*(?:[-•*]|\d+[.)-]?)\s*",
            "",
            linea,
        )
        linea_limpia = limpiar_texto(linea)

        if linea_limpia:
            actividades.append(linea_limpia)

    if len(actividades) <= 1 and contenido.count(",") >= 2:
        actividades = [
            limpiar_texto(item)
            for item in contenido.split(",")
            if limpiar_texto(item)
        ]

    return actividades


def limpiar_parrafo(parrafo: Paragraph) -> None:
    for elemento in list(parrafo._p):
        if elemento.tag != qn("w:pPr"):
            parrafo._p.remove(elemento)


def escribir_parrafo(
    parrafo: Paragraph,
    texto: str,
    *,
    tamano: float = 11,
    negrita: bool = False,
    alineacion=None,
) -> None:
    limpiar_parrafo(parrafo)
    run = parrafo.add_run(str(texto or ""))
    run.bold = negrita
    run.font.name = "Arial"
    run.font.size = Pt(tamano)

    if alineacion is not None:
        parrafo.alignment = alineacion


def escribir_celda(celda, texto: object) -> None:
    celda.text = ""
    parrafo = celda.paragraphs[0]
    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = parrafo.add_run(str(texto or ""))
    run.font.name = "Arial"
    run.font.size = Pt(10)


def texto_elemento_xml(elemento) -> str:
    return "".join(
        nodo.text or ""
        for nodo in elemento.xpath('.//*[local-name()="t"]')
    ).strip()


def buscar_parrafo(
    documento: Document,
    texto_objetivo: str,
    *,
    coincidencia_exacta: bool = False,
) -> Paragraph:
    objetivo = " ".join(str(texto_objetivo).split()).casefold()

    for parrafo in documento.paragraphs:
        texto = " ".join(parrafo.text.split()).casefold()

        if coincidencia_exacta and texto == objetivo:
            return parrafo

        if not coincidencia_exacta and objetivo in texto:
            return parrafo

    raise ValueError(
        f"No se encontró el apartado '{texto_objetivo}' en la plantilla."
    )


def parrafo_siguiente(parrafo: Paragraph) -> Paragraph:
    elemento = parrafo._p.getnext()

    while elemento is not None:
        if elemento.tag == qn("w:p"):
            return Paragraph(elemento, parrafo._parent)
        elemento = elemento.getnext()

    raise ValueError(
        f"No se encontró un párrafo después de '{parrafo.text}'."
    )


def insertar_parrafo_despues(
    parrafo: Paragraph,
    texto: str = "",
    estilo: str | None = None,
) -> Paragraph:
    nuevo_p = OxmlElement("w:p")
    parrafo._p.addnext(nuevo_p)
    nuevo = Paragraph(nuevo_p, parrafo._parent)

    if estilo:
        try:
            nuevo.style = estilo
        except KeyError:
            pass

    if texto:
        escribir_parrafo(nuevo, texto)

    return nuevo


def eliminar_instrucciones_y_control_cambios(documento: Document) -> None:
    """Elimina desde INSTRUCCIONES hasta antes de las propiedades de sección."""
    cuerpo = documento._element.body
    eliminar = False

    for elemento in list(cuerpo):
        if elemento.tag == qn("w:sectPr"):
            continue

        texto = texto_elemento_xml(elemento)

        if (
            elemento.tag == qn("w:p")
            and texto.strip().casefold().startswith("instrucciones")
        ):
            eliminar = True

        if eliminar:
            cuerpo.remove(elemento)


def marcar_actualizacion_campos(documento: Document) -> None:
    settings = documento.settings.element
    update_fields = settings.find(qn("w:updateFields"))

    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)

    update_fields.set(qn("w:val"), "true")

    for campo in documento._element.xpath(
        './/*[local-name()="fldChar" and @w:fldCharType="begin"]'
    ):
        campo.set(qn("w:dirty"), "true")


def marcar_clasificacion(documento: Document, clasificacion: str) -> None:
    """
    Escribe una X dentro del cuadro de texto de la clasificación seleccionada,
    conservando los cuadros y la diagramación original de la portada.
    """
    tabla_portada = documento.tables[0]
    fila = tabla_portada.rows[5]

    for celda in fila.cells:
        nombre = " ".join(celda.text.split())
        seleccionada = nombre.casefold() == clasificacion.casefold()

        contenidos = celda._tc.xpath(
            './/*[local-name()="txbxContent"]'
        )

        for contenido in contenidos:
            parrafos = contenido.xpath('./*[local-name()="p"]')

            if not parrafos:
                nuevo_p = OxmlElement("w:p")
                contenido.append(nuevo_p)
                parrafos = [nuevo_p]

            for p_xml in parrafos:
                for hijo in list(p_xml):
                    if hijo.tag != qn("w:pPr"):
                        p_xml.remove(hijo)

                p_pr = p_xml.find(qn("w:pPr"))
                if p_pr is None:
                    p_pr = OxmlElement("w:pPr")
                    p_xml.insert(0, p_pr)

                jc = p_pr.find(qn("w:jc"))
                if jc is None:
                    jc = OxmlElement("w:jc")
                    p_pr.append(jc)
                jc.set(qn("w:val"), "center")

                if seleccionada:
                    run = OxmlElement("w:r")
                    r_pr = OxmlElement("w:rPr")

                    negrita = OxmlElement("w:b")
                    r_pr.append(negrita)

                    tamano = OxmlElement("w:sz")
                    tamano.set(qn("w:val"), "20")
                    r_pr.append(tamano)

                    run.append(r_pr)
                    texto = OxmlElement("w:t")
                    texto.text = "X"
                    run.append(texto)
                    p_xml.append(run)


def llenar_tabla_informacion_general(
    documento: Document,
    datos: dict,
) -> None:
    tabla = documento.tables[1]

    valores = [
        datos.get("nombre_talento", ""),
        datos.get("nombre_proyecto", ""),
        datos.get("codigo_proyecto", ""),
        datos.get("nombre_experto", ""),
        datos.get("linea_tecnologica", ""),
        datos.get("trl_inicial", ""),
        datos.get("trl_alcanzado", ""),
        datos.get("tecnoparque", ""),
        datos.get("fecha_entrega_texto", ""),
    ]

    for indice, valor in enumerate(valores):
        escribir_celda(tabla.rows[indice].cells[1], valor)


def formatear_lista_documento(
    elementos: list[str],
    *,
    prefijo: str = "•",
) -> str:
    return "\n".join(
        f"{prefijo} {elemento}"
        for elemento in elementos
        if limpiar_texto(elemento)
    )


def escribir_lista_en_parrafos(
    parrafo_inicial: Paragraph,
    elementos: list[str],
    *,
    prefijo: str = "•",
) -> Paragraph:
    """Escribe cada elemento en un párrafo independiente y retorna el último."""
    elementos_limpios = [
        limpiar_texto(str(elemento))
        for elemento in elementos
        if limpiar_texto(str(elemento))
    ]

    if not elementos_limpios:
        elementos_limpios = ["No se suministró información para este apartado."]

    actual = parrafo_inicial

    for indice, elemento in enumerate(elementos_limpios):
        if indice > 0:
            actual = insertar_parrafo_despues(
                actual,
                estilo="Normal",
            )

        escribir_parrafo(
            actual,
            f"{prefijo} {elemento}",
            tamano=11,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
        )
        actual.paragraph_format.space_after = Pt(4)
        actual.paragraph_format.left_indent = Cm(0.35)
        actual.paragraph_format.first_line_indent = Cm(-0.35)

    return actual


def eliminar_parrafo(parrafo: Paragraph) -> None:
    elemento = parrafo._element
    elemento.getparent().remove(elemento)
    parrafo._p = None
    parrafo._element = None


def configurar_estilos_y_tabla_contenido(documento: Document) -> None:
    """
    Garantiza que los títulos del informe usen estilos de Word y que la tabla
    de contenido se actualice con la paginación real al abrir el archivo.
    """
    titulos_nivel_1 = [
        "Información general del proyecto",
        "Introducción",
        "Planteamiento del problema",
        "Objetivos",
        "5. Estado del arte y estado de la técnica",
        "6. Metodología de desarrollo",
        "7. Desarrollo del proyecto",
        "8. Normatividad",
        "9. Resultados obtenidos",
        "10. Análisis de viabilidad",
        "11. Propiedad intelectual y transferencia tecnológica",
        "12. Impacto del proyecto",
        "13. Conclusiones",
        "14. Referencias bibliográficas",
        "15. Anexos",
    ]
    titulos_nivel_2 = [
        "4.1 Objetivo General",
        "4.2 Objetivos Específicos",
    ]

    for titulo in titulos_nivel_1:
        buscar_parrafo(
            documento,
            titulo,
            coincidencia_exacta=True,
        ).style = "Heading 1"

    for titulo in titulos_nivel_2:
        buscar_parrafo(
            documento,
            titulo,
            coincidencia_exacta=True,
        ).style = "Heading 2"

    # La plantilla ya contiene un campo TOC. Se limita a niveles 1 y 2,
    # se desbloquea y se marca como pendiente de actualización.
    for instr in documento._element.xpath(
        './/*[local-name()="instrText"]'
    ):
        if "TOC" in (instr.text or ""):
            instr.text = r' TOC \o "1-2" \h \z \u '

    # Elimina marcadores internos obsoletos del índice para que Word o
    # LibreOffice los reconstruyan según la ubicación actual de cada título.
    ids_toc: set[str] = set()

    for marcador in documento._element.xpath(
        './/*[local-name()="bookmarkStart"]'
    ):
        nombre = marcador.get(qn("w:name"), "")
        if nombre.startswith("_Toc"):
            identificador = marcador.get(qn("w:id"))
            if identificador is not None:
                ids_toc.add(identificador)
            marcador.getparent().remove(marcador)

    for marcador in documento._element.xpath(
        './/*[local-name()="bookmarkEnd"]'
    ):
        if marcador.get(qn("w:id")) in ids_toc:
            marcador.getparent().remove(marcador)

    for campo in documento._element.xpath(
        './/*[local-name()="fldChar"]'
    ):
        campo.attrib.pop(qn("w:fldLock"), None)
        if campo.get(qn("w:fldCharType")) == "begin":
            campo.set(qn("w:dirty"), "true")

    marcar_actualizacion_campos(documento)


def insertar_desarrollo_actividades(
    parrafo_ancla: Paragraph,
    actividades: list[dict],
) -> Paragraph:
    """Inserta una subsección por actividad y deja espacio para evidencias."""
    cursor = parrafo_ancla

    for numero, actividad in enumerate(actividades, start=1):
        titulo = limpiar_frases_meta(str(actividad.get("titulo", "")))
        fase = limpiar_frases_meta(
            str(actividad.get("fase_metodologica", ""))
        )
        descripcion = limpiar_frases_meta(
            str(actividad.get("descripcion_tecnica", ""))
        )

        cursor = insertar_parrafo_despues(cursor)
        escribir_parrafo(
            cursor,
            f"Actividad {numero}. {titulo or f'Actividad técnica {numero}'}",
            tamano=11,
            negrita=True,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
        )
        cursor.paragraph_format.space_before = Pt(8)
        cursor.paragraph_format.space_after = Pt(4)

        cursor = insertar_parrafo_despues(cursor)
        escribir_parrafo(
            cursor,
            f"Fase metodológica relacionada: {fase}",
            tamano=10.5,
            negrita=True,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
        )
        cursor.paragraph_format.space_after = Pt(4)

        cursor = insertar_parrafo_despues(cursor)
        escribir_parrafo(
            cursor,
            descripcion,
            tamano=11,
            alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        cursor.paragraph_format.space_after = Pt(5)

        cursor = insertar_parrafo_despues(cursor)
        escribir_parrafo(
            cursor,
            "Evidencia de la actividad:",
            tamano=11,
            negrita=True,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
        )

        cursor = insertar_parrafo_despues(cursor)
        escribir_parrafo(
            cursor,
            (
                "[Insertar fotografías, capturas de pantalla, enlaces o soportes "
                "correspondientes a esta actividad]"
            ),
            tamano=10.5,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
        )
        cursor.paragraph_format.space_after = Pt(18)

    return cursor


def renumerar_apartados_posteriores(documento: Document) -> None:
    """Desplaza la numeración de los apartados posteriores a Desarrollo."""
    cambios = [
        ("14. Anexos", "15. Anexos"),
        ("13. Referencias bibliográficas", "14. Referencias bibliográficas"),
        ("12. Conclusiones", "13. Conclusiones"),
        ("11. Impacto del proyecto", "12. Impacto del proyecto"),
        (
            "10. Propiedad intelectual y transferencia tecnológica",
            "11. Propiedad intelectual y transferencia tecnológica",
        ),
        ("9. Análisis de viabilidad", "10. Análisis de viabilidad"),
        ("8. Resultados obtenidos", "9. Resultados obtenidos"),
    ]

    for titulo_actual, titulo_nuevo in cambios:
        parrafo = buscar_parrafo(
            documento,
            titulo_actual,
            coincidencia_exacta=True,
        )
        escribir_parrafo(
            parrafo,
            titulo_nuevo,
            tamano=11,
            negrita=True,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
        )


def insertar_normatividad(
    parrafo_ancla: Paragraph,
    texto_normativo: str,
    normas: list[dict],
) -> Paragraph:
    """Inserta el apartado Normatividad y lista sus fuentes aplicables."""
    encabezado = insertar_parrafo_despues(
        parrafo_ancla,
        "8. Normatividad",
        estilo="Heading 1",
    )
    escribir_parrafo(
        encabezado,
        "8. Normatividad",
        tamano=11,
        negrita=True,
        alineacion=WD_ALIGN_PARAGRAPH.LEFT,
    )

    cursor = insertar_parrafo_despues(encabezado)
    escribir_parrafo(
        cursor,
        limpiar_texto(texto_normativo),
        tamano=11,
        alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    cursor.paragraph_format.space_after = Pt(8)

    normas_normalizadas = normalizar_normas_aplicables(normas)

    for ambito, subtitulo in [
        ("Colombia", "Normatividad colombiana"),
        ("Internacional", "Normatividad y estándares internacionales"),
    ]:
        normas_ambito = [
            item
            for item in normas_normalizadas
            if item["ambito"] == ambito
        ]

        if not normas_ambito:
            continue

        cursor = insertar_parrafo_despues(cursor)
        escribir_parrafo(
            cursor,
            subtitulo,
            tamano=11,
            negrita=True,
            alineacion=WD_ALIGN_PARAGRAPH.LEFT,
        )
        cursor.paragraph_format.space_before = Pt(6)
        cursor.paragraph_format.space_after = Pt(4)

        for item in normas_ambito:
            cursor = insertar_parrafo_despues(cursor)
            texto_item = (
                f"• {item['norma']} — {item['entidad']}. "
                f"{item['aplicacion']} "
                f"Carácter de aplicación: {item['caracter_aplicacion']}. "
                f"Fuente oficial: {item['fuente_oficial']}"
            )
            escribir_parrafo(
                cursor,
                texto_item,
                tamano=10.5,
                alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY,
            )
            cursor.paragraph_format.left_indent = Cm(0.35)
            cursor.paragraph_format.first_line_indent = Cm(-0.35)
            cursor.paragraph_format.space_after = Pt(5)

    return cursor


def insertar_tabla_resultados(
    documento: Document,
    parrafo_ancla: Paragraph,
    entregables: list[str],
) -> None:
    entregables_limpios = [
        limpiar_texto(str(entregable))
        for entregable in entregables
        if limpiar_texto(str(entregable))
    ]

    if not entregables_limpios:
        entregables_limpios = [
            "No se registraron entregables obtenidos para relacionar."
        ]

    tabla = documento.add_table(
        rows=1,
        cols=3,
    )
    tabla.style = "Table Grid"
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.autofit = False

    anchos = [Cm(1.2), Cm(11.3), Cm(4.6)]
    encabezados = [
        "N.°",
        "Descripción del entregable obtenido",
        "Evidencia",
    ]

    for indice, encabezado in enumerate(encabezados):
        celda = tabla.rows[0].cells[indice]
        celda.width = anchos[indice]
        celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        escribir_celda(celda, encabezado)
        for run in celda.paragraphs[0].runs:
            run.bold = True
        celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for numero, entregable in enumerate(entregables_limpios, start=1):
        fila = tabla.add_row()
        valores = [
            str(numero),
            entregable,
            "Agregar enlace: ______________________________",
        ]

        for indice, valor in enumerate(valores):
            celda = fila.cells[indice]
            celda.width = anchos[indice]
            celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            escribir_celda(celda, valor)

        fila.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    parrafo_ancla._p.addnext(tabla._tbl)



def serializar_datos_informe(datos: dict) -> dict:
    resultado = deepcopy(datos)

    fecha_entrega = resultado.get("fecha_entrega")
    if isinstance(fecha_entrega, date):
        resultado["fecha_entrega"] = fecha_entrega.strftime("%d/%m/%Y")

    return resultado


# =====================================================
# GENERACIÓN DE CONTENIDO
# =====================================================

def inferir_metodologia_base(datos: dict) -> str:
    """Infiere una metodología coherente sin solicitarla en el formulario."""
    texto_base = " ".join(
        [
            datos.get("descripcion_general_proyecto", ""),
            datos.get("entregables_proyecto_base", ""),
            datos.get("innovacion_proyecto_base", ""),
            actividades_en_texto(
                datos.get("actividades_ejecutadas_base", [])
            ),
        ]
    ).casefold()

    if any(
        termino in texto_base
        for termino in [
            "inteligencia artificial",
            "machine learning",
            "modelo predictivo",
            "analítica de datos",
            "clasificación",
            "dataset",
        ]
    ):
        return (
            "CRISP-DM, complementada con desarrollo iterativo y validación "
            "experimental"
        )

    if any(
        termino in texto_base
        for termino in [
            "software",
            "aplicación",
            "plataforma",
            "sistema web",
            "app",
            "interfaz",
        ]
    ):
        return (
            "Diseño Centrado en el Usuario, Design Thinking y desarrollo "
            "ágil iterativo"
        )

    if any(
        termino in texto_base
        for termino in [
            "electrónica",
            "sensor",
            "microcontrolador",
            "iot",
            "automatización",
            "circuito",
        ]
    ):
        return (
            "ingeniería de sistemas mediante modelo V, complementada con "
            "prototipado iterativo"
        )

    if any(
        termino in texto_base
        for termino in [
            "impresión 3d",
            "manufactura",
            "mecanismo",
            "carcasa",
            "pieza",
            "producto físico",
            "diseño industrial",
        ]
    ):
        return (
            "Design Thinking, desarrollo iterativo de prototipos y Diseño "
            "para Manufactura y Ensamble (DFMA)"
        )

    if any(
        termino in texto_base
        for termino in [
            "marca",
            "identidad visual",
            "experiencia de usuario",
            "servicio",
            "comunicación visual",
        ]
    ):
        return "Doble Diamante y Diseño Centrado en el Usuario"

    if any(
        termino in texto_base
        for termino in [
            "agroindustrial",
            "alimento",
            "formulación",
            "proceso productivo",
            "laboratorio",
        ]
    ):
        return (
            "investigación aplicada, validación experimental y mejora "
            "continua mediante DMAIC"
        )

    return "Design Thinking y desarrollo iterativo de prototipos"


def metodologias_en_texto(datos: dict) -> str:
    seleccionadas = datos.get("metodologias_seleccionadas", []) or []
    otra = limpiar_texto(datos.get("otra_metodologia", ""))

    metodologias = [
        limpiar_texto(str(item))
        for item in seleccionadas
        if limpiar_texto(str(item)) and str(item) != "Otra"
    ]

    if otra:
        metodologias.append(otra)

    if metodologias:
        return ", ".join(metodologias)

    metodologia_respaldo = limpiar_texto(
        datos.get("metodologia_inferida", "")
    )
    return metodologia_respaldo or inferir_metodologia_base(datos)


def fases_metodologia(nombre: str) -> str:
    """Devuelve las etapas estándar de una metodología seleccionada."""
    clave = limpiar_texto(nombre).casefold()

    catalogo = {
        "metodologías ágiles": (
            "priorización de requerimientos, planificación iterativa, ejecución "
            "en ciclos cortos, revisión de avances, retroalimentación y mejora continua"
        ),
        "scrum": (
            "definición y priorización del Product Backlog, planificación del Sprint, "
            "ejecución iterativa, seguimiento diario, Sprint Review y retrospectiva"
        ),
        "kanban": (
            "visualización del flujo de trabajo, limitación del trabajo en curso, "
            "gestión del flujo, seguimiento de indicadores y mejora continua"
        ),
        "modelo en cascada": (
            "levantamiento y análisis de requisitos, diseño, implementación, pruebas, "
            "despliegue y mantenimiento"
        ),
        "modelo espiral": (
            "definición de objetivos, análisis de riesgos, desarrollo o prototipado, "
            "evaluación de resultados y planificación del siguiente ciclo"
        ),
        "design thinking": "empatizar, definir, idear, prototipar y evaluar",
        "doble diamante": "descubrir, definir, desarrollar y entregar",
        "diseño centrado en el usuario (dcu)": (
            "comprensión del contexto de uso, especificación de requerimientos, "
            "diseño de soluciones, evaluación con usuarios e iteración"
        ),
        "diseño para manufactura y ensamble (dfma)": (
            "definición de requerimientos, simplificación del diseño, análisis de "
            "manufacturabilidad, análisis de ensamblaje y validación del producto"
        ),
        "lean startup": (
            "formular hipótesis, construir el producto mínimo viable, medir resultados, "
            "aprender de la evidencia y decidir entre perseverar o ajustar"
        ),
        "stage-gate": (
            "descubrimiento, definición de alcance, formulación del caso, desarrollo, "
            "pruebas, validación y decisión de avance mediante puntos de control"
        ),
        "desarrollo iterativo de prototipos": (
            "definición de requerimientos, construcción del prototipo, prueba, "
            "retroalimentación, ajuste y nueva iteración"
        ),
        "ingeniería de sistemas y modelo v": (
            "requisitos, diseño del sistema, diseño detallado, implementación, "
            "verificación unitaria, integración, validación del sistema y aceptación"
        ),
        "crisp-dm para proyectos de datos e inteligencia artificial": (
            "comprensión del negocio, comprensión de los datos, preparación de datos, "
            "modelado, evaluación y despliegue"
        ),
        "dmaic / six sigma": "definir, medir, analizar, mejorar y controlar",
        "investigación aplicada y validación experimental": (
            "formulación del problema, diseño metodológico, experimentación, análisis "
            "de resultados, validación y documentación"
        ),
        "triz para solución inventiva de problemas": (
            "definición del problema, identificación de contradicciones, selección de "
            "principios inventivos, formulación de alternativas y evaluación técnica"
        ),
    }

    return catalogo.get(
        clave,
        (
            "definición del alcance, planeación, ejecución, seguimiento, "
            "validación y cierre conforme al enfoque seleccionado"
        ),
    )


def fases_metodologias_en_texto(datos: dict) -> str:
    """Presenta cada metodología seleccionada con sus fases de referencia."""
    seleccionadas = datos.get("metodologias_seleccionadas", []) or []
    otra = limpiar_texto(datos.get("otra_metodologia", ""))

    metodologias = [
        limpiar_texto(str(item))
        for item in seleccionadas
        if limpiar_texto(str(item)) and str(item) != "Otra"
    ]

    if otra:
        metodologias.append(otra)

    if not metodologias:
        metodologias = [metodologias_en_texto(datos)]

    return "\n".join(
        f"- {metodologia}: {fases_metodologia(metodologia)}."
        for metodologia in metodologias
    )


def limpiar_frases_meta(texto: str) -> str:
    """Elimina explicaciones sobre el proceso de generación o corrección."""
    contenido = limpiar_texto(texto)
    if not contenido:
        return ""

    frases_prohibidas = (
        "fueron previamente corregid",
        "previamente corregid",
        "sin copiar literalmente",
        "no se copiaron",
        "texto ingresado",
        "texto diligenciado",
        "generado por inteligencia artificial",
        "generado por ia",
        "prompt",
        "antes de ser incorporad",
        "en lugar de reproducir",
        "reformuladas antes de",
    )

    oraciones = re.split(r"(?<=[.!?])\s+", contenido)
    conservadas = [
        oracion.strip()
        for oracion in oraciones
        if oracion.strip()
        and not any(frase in oracion.casefold() for frase in frases_prohibidas)
    ]
    return limpiar_texto(" ".join(conservadas))


def metodologia_modo_prueba(datos: dict) -> str:
    metodologias = metodologias_en_texto(datos)
    fases = fases_metodologias_en_texto(datos).replace("\n", " ")
    return (
        f"El proyecto se desarrolló mediante {metodologias}. La aplicación del "
        f"enfoque consideró las siguientes fases de referencia: {fases} Estas "
        "etapas permitieron ordenar el trabajo desde la comprensión de la necesidad "
        "y la definición de requerimientos hasta el diseño, la implementación, la "
        "revisión y la validación de la solución. Cuando se emplearon varios enfoques, "
        "su integración se realizó de manera complementaria, asignando a cada uno una "
        "función concreta dentro del proceso. Las metodologías orientadas al usuario "
        "aportaron criterios para comprender el contexto, definir necesidades y revisar "
        "la pertinencia de las decisiones. Los enfoques iterativos facilitaron la "
        "priorización, el seguimiento de avances y la incorporación progresiva de "
        "ajustes. Los modelos de ingeniería permitieron mantener trazabilidad entre "
        "requerimientos, diseño, construcción y verificación."
    )


def desarrollo_actividades_modo_prueba(
    datos: dict,
    actividades_corregidas: list[str],
) -> list[dict]:
    metodologia = metodologias_en_texto(datos)
    resultado: list[dict] = []
    for indice, actividad in enumerate(actividades_corregidas, start=1):
        titulo = limpiar_texto(actividad).rstrip(".")
        resultado.append(
            {
                "titulo": titulo or f"Actividad técnica {indice}",
                "fase_metodologica": metodologia,
                "descripcion_tecnica": (
                    f"La actividad se ejecutó en correspondencia con {metodologia}, "
                    "aplicando procedimientos y decisiones vinculados con el alcance "
                    "técnico definido. Su desarrollo mantuvo relación con los objetivos, "
                    "los componentes de la solución y la secuencia metodológica del "
                    "proyecto. La revisión de avances permitió identificar condiciones "
                    "de implementación, necesidades de integración y aspectos sujetos a "
                    "ajuste, sin atribuir resultados no documentados."
                ),
                "evidencia": (
                    "Insertar fotografías, capturas de pantalla, enlaces o soportes "
                    "correspondientes a esta actividad."
                ),
            }
        )
    return resultado


def normalizar_desarrollo_actividades(
    valor: object,
    actividades_corregidas: list[str],
    datos: dict,
) -> list[dict]:
    elementos = valor if isinstance(valor, list) else []
    respaldo = desarrollo_actividades_modo_prueba(datos, actividades_corregidas)
    resultado: list[dict] = []

    for indice in range(len(actividades_corregidas)):
        item = elementos[indice] if indice < len(elementos) else {}
        if not isinstance(item, dict):
            item = {}
        base = respaldo[indice]

        resultado.append(
            {
                "titulo": limpiar_frases_meta(str(item.get("titulo", "")))
                or base["titulo"],
                "fase_metodologica": limpiar_frases_meta(
                    str(item.get("fase_metodologica", ""))
                )
                or base["fase_metodologica"],
                "descripcion_tecnica": limpiar_frases_meta(
                    str(item.get("descripcion_tecnica", ""))
                )
                or base["descripcion_tecnica"],
                "evidencia": (
                    "Insertar fotografías, capturas de pantalla, enlaces o soportes "
                    "correspondientes a esta actividad."
                ),
            }
        )
    return resultado


def actividades_en_texto(actividades: list[str]) -> str:
    actividades_limpias = [
        limpiar_texto(str(actividad))
        for actividad in actividades
        if limpiar_texto(str(actividad))
    ]

    return "\n".join(
        f"{indice}. {actividad}"
        for indice, actividad in enumerate(actividades_limpias, start=1)
    )


def corregir_entregable_basico(texto: str) -> str:
    """
    Corrección de respaldo para modo prueba.

    La mejora técnica avanzada se realiza con la API; esta función conserva el
    sentido original, normaliza la redacción y evita inventar características.
    """
    entregable = limpiar_texto(texto)

    if not entregable:
        return ""

    entregable = re.sub(
        r"^\s*(?:[-•*]|\d+[.)-]?)\s*",
        "",
        entregable,
    ).strip()

    if not entregable:
        return ""

    entregable = entregable[0].upper() + entregable[1:]

    if entregable[-1] not in ".!?":
        entregable += "."

    return entregable


def corregir_actividad_basica(texto: str) -> str:
    actividad = limpiar_texto(texto)

    if not actividad:
        return ""

    actividad = actividad[0].upper() + actividad[1:]

    if actividad[-1] not in ".!?":
        actividad += "."

    return actividad


def obtener_api_key() -> str | None:
    api_key = os.getenv("OPENAI_API_KEY")

    if api_key:
        return api_key.strip()

    try:
        valor = st.secrets.get("OPENAI_API_KEY")
        return str(valor).strip() if valor else None
    except Exception:
        return None


def extraer_json_respuesta(texto: str) -> dict:
    contenido = str(texto or "").strip()
    contenido = re.sub(r"^```(?:json)?\s*", "", contenido, flags=re.I)
    contenido = re.sub(r"\s*```$", "", contenido)

    try:
        resultado = json.loads(contenido)
        return resultado if isinstance(resultado, dict) else {}
    except json.JSONDecodeError:
        coincidencia = re.search(r"\{.*\}", contenido, flags=re.S)

        if not coincidencia:
            return {}

        try:
            resultado = json.loads(coincidencia.group(0))
            return resultado if isinstance(resultado, dict) else {}
        except json.JSONDecodeError:
            return {}


def referentes_modo_prueba(datos: dict) -> list[dict]:
    """
    Referentes reales y generales para validar la estructura en modo prueba.
    En modo API se reemplazan por soluciones específicamente relacionadas.
    """
    descripcion = datos.get("descripcion_general_proyecto", "").casefold()

    if any(
        palabra in descripcion
        for palabra in ["inteligencia artificial", "modelo", "clasificación", "machine learning"]
    ):
        return [
            {
                "nombre": "Teachable Machine",
                "entidad": "Google Creative Lab",
                "anio": "s. f.",
                "descripcion": (
                    "Herramienta web que permite crear modelos de aprendizaje automático "
                    "con imágenes, sonidos y posturas sin requerir programación avanzada."
                ),
                "cita_corta": "(Google Creative Lab, s. f.)",
                "referencia_apa": (
                    "Google Creative Lab. (s. f.). Teachable Machine. "
                    "https://teachablemachine.withgoogle.com/"
                ),
                "url": "https://teachablemachine.withgoogle.com/",
            },
            {
                "nombre": "TensorFlow Lite",
                "entidad": "Google",
                "anio": "s. f.",
                "descripcion": (
                    "Conjunto de herramientas para ejecutar modelos de aprendizaje "
                    "automático en dispositivos móviles, embebidos y de borde."
                ),
                "cita_corta": "(Google, s. f.)",
                "referencia_apa": (
                    "Google. (s. f.). TensorFlow Lite. "
                    "https://www.tensorflow.org/lite"
                ),
                "url": "https://www.tensorflow.org/lite",
            },
        ]

    return [
        {
            "nombre": "MIT App Inventor",
            "entidad": "Massachusetts Institute of Technology",
            "anio": "s. f.",
            "descripcion": (
                "Entorno de desarrollo visual orientado a la creación de aplicaciones "
                "móviles mediante bloques y procesos iterativos de prototipado."
            ),
            "cita_corta": "(MIT App Inventor, s. f.)",
            "referencia_apa": (
                "MIT App Inventor. (s. f.). MIT App Inventor. "
                "https://appinventor.mit.edu/"
            ),
            "url": "https://appinventor.mit.edu/",
        },
        {
            "nombre": "Arduino",
            "entidad": "Arduino",
            "anio": "s. f.",
            "descripcion": (
                "Plataforma abierta de hardware y software utilizada para construir "
                "prototipos electrónicos interactivos y soluciones conectadas."
            ),
            "cita_corta": "(Arduino, s. f.)",
            "referencia_apa": (
                "Arduino. (s. f.). Arduino documentation. "
                "https://docs.arduino.cc/"
            ),
            "url": "https://docs.arduino.cc/",
        },
    ]


def investigar_referentes_reales(
    datos: dict,
    modelo_openai: str,
) -> list[dict]:
    """
    Busca exactamente dos proyectos o soluciones reales y devuelve referencias
    verificables. La búsqueda usa la herramienta web de la Responses API.
    """
    if OpenAI is None:
        raise RuntimeError(
            "La librería openai no está instalada y no es posible verificar "
            "los dos referentes reales del Estado del Arte."
        )

    api_key = obtener_api_key()

    if not api_key:
        raise RuntimeError(
            "No se encontró OPENAI_API_KEY para investigar los referentes reales."
        )

    client = OpenAI(api_key=api_key)

    prompt = f"""
Busca en fuentes oficiales, institucionales o académicas exactamente dos proyectos,
productos o soluciones reales comparables con el siguiente proyecto de base tecnológica.

Proyecto: {datos.get('nombre_proyecto', '')}
Descripción: {datos.get('descripcion_general_proyecto', '')}
Línea tecnológica: {datos.get('linea_tecnologica', '')}

Selecciona referentes realmente relacionados con el propósito, la tecnología o el
problema atendido. No inventes nombres, entidades, años, enlaces ni publicaciones.
Prefiere la página oficial del proyecto o una publicación académica primaria.

Responde exclusivamente con JSON válido:
{{
  "proyectos": [
    {{
      "nombre": "nombre oficial",
      "entidad": "entidad responsable",
      "anio": "año o s. f.",
      "descripcion": "síntesis verificable de 50 a 90 palabras",
      "cita_corta": "(Entidad, año)",
      "referencia_apa": "referencia completa en APA 7 con URL",
      "url": "https://..."
    }}
  ]
}}
"""

    respuesta = client.responses.create(
        model=modelo_openai,
        tools=[{"type": "web_search"}],
        input=prompt,
        temperature=0.1,
    )

    datos_respuesta = extraer_json_respuesta(
        getattr(respuesta, "output_text", "")
    )
    proyectos = datos_respuesta.get("proyectos", [])

    if not isinstance(proyectos, list):
        proyectos = []

    proyectos_validos: list[dict] = []

    for item in proyectos:
        if not isinstance(item, dict):
            continue

        campos = {
            clave: limpiar_texto(str(item.get(clave, "")))
            for clave in [
                "nombre",
                "entidad",
                "anio",
                "descripcion",
                "cita_corta",
                "referencia_apa",
                "url",
            ]
        }

        if (
            campos["nombre"]
            and campos["entidad"]
            and campos["descripcion"]
            and campos["referencia_apa"]
            and campos["url"].startswith(("http://", "https://"))
        ):
            proyectos_validos.append(campos)

    if len(proyectos_validos) < 2:
        raise RuntimeError(
            "La búsqueda no produjo dos referentes reales con fuentes verificables. "
            "Intenta nuevamente o revisa la descripción general del proyecto."
        )

    return proyectos_validos[:2]


def es_fuente_normativa_oficial(url: str) -> bool:
    """Valida que la fuente pertenezca a una entidad oficial o normalizadora."""
    try:
        dominio = urlparse(str(url or "")).netloc.casefold()
    except Exception:
        return False

    if dominio.startswith("www."):
        dominio = dominio[4:]

    dominios_permitidos = (
        "suin-juriscol.gov.co",
        "funcionpublica.gov.co",
        "sic.gov.co",
        "mintic.gov.co",
        "mincit.gov.co",
        "minenergia.gov.co",
        "minambiente.gov.co",
        "minsalud.gov.co",
        "invima.gov.co",
        "mintrabajo.gov.co",
        "derechodeautor.gov.co",
        "comunidadandina.org",
        "icontec.org",
        "iso.org",
        "iec.ch",
        "w3.org",
        "itu.int",
        "who.int",
        "fao.org",
        "codexalimentarius.fao.org",
        "oecd.org",
        "un.org",
    )

    return (
        dominio.endswith(".gov.co")
        or any(
            dominio == permitido or dominio.endswith(f".{permitido}")
            for permitido in dominios_permitidos
        )
        or dominio.endswith(".europa.eu")
        or dominio == "europa.eu"
        or dominio.endswith(".ieee.org")
        or dominio == "ieee.org"
        or dominio.endswith(".unesco.org")
        or dominio == "unesco.org"
    )


def normalizar_normas_aplicables(valor: object) -> list[dict]:
    """Normaliza y valida la lista de normas identificadas para el proyecto."""
    elementos = valor if isinstance(valor, list) else []
    resultado: list[dict] = []

    for item in elementos:
        if not isinstance(item, dict):
            continue

        norma = limpiar_texto(str(item.get("norma", "")))
        ambito = limpiar_texto(str(item.get("ambito", "")))
        entidad = limpiar_texto(str(item.get("entidad", "")))
        aplicacion = limpiar_texto(str(item.get("aplicacion", "")))
        caracter = limpiar_texto(
            str(item.get("caracter_aplicacion", ""))
        )
        fuente = str(item.get("fuente_oficial", "")).strip()

        if ambito.casefold().startswith("col"):
            ambito = "Colombia"
        elif ambito:
            ambito = "Internacional"

        if (
            norma
            and ambito
            and entidad
            and aplicacion
            and caracter
            and fuente.startswith(("http://", "https://"))
            and es_fuente_normativa_oficial(fuente)
        ):
            resultado.append(
                {
                    "ambito": ambito,
                    "norma": norma,
                    "entidad": entidad,
                    "aplicacion": aplicacion,
                    "caracter_aplicacion": caracter,
                    "fuente_oficial": fuente,
                }
            )

    return resultado[:10]


def normatividad_modo_prueba(datos: dict) -> dict:
    """
    Genera una selección conservadora de normas para validar la estructura local.

    En modo API esta información se reemplaza por una búsqueda web específica
    para el proyecto.
    """
    texto_base = " ".join(
        [
            datos.get("nombre_proyecto", ""),
            datos.get("descripcion_general_proyecto", ""),
            datos.get("entregables_proyecto_base", ""),
            datos.get("innovacion_proyecto_base", ""),
            actividades_en_texto(
                datos.get("actividades_ejecutadas_base", [])
            ),
        ]
    ).casefold()

    normas: list[dict] = []

    if any(
        termino in texto_base
        for termino in [
            "dato",
            "usuario",
            "registro",
            "plataforma",
            "aplicación",
            "sistema web",
            "base de datos",
            "formulario",
            "cámara",
            "biométr",
        ]
    ):
        normas.extend(
            [
                {
                    "ambito": "Colombia",
                    "norma": "Ley 1581 de 2012",
                    "entidad": "Congreso de Colombia",
                    "aplicacion": (
                        "Debe considerarse cuando la solución recolecte, almacene, "
                        "consulte, transmita o elimine datos personales."
                    ),
                    "caracter_aplicacion": (
                        "Obligatoria cuando exista tratamiento de datos personales"
                    ),
                    "fuente_oficial": (
                        "https://www.suin-juriscol.gov.co/"
                        "viewDocument.asp?id=1684507"
                    ),
                },
                {
                    "ambito": "Internacional",
                    "norma": "ISO/IEC 27001:2022",
                    "entidad": (
                        "International Organization for Standardization e IEC"
                    ),
                    "aplicacion": (
                        "Sirve como referencia para gestionar riesgos de seguridad "
                        "de la información, controles, acceso y mejora continua."
                    ),
                    "caracter_aplicacion": (
                        "Referencia técnica voluntaria, salvo exigencia contractual"
                    ),
                    "fuente_oficial": "https://www.iso.org/standard/27001",
                },
            ]
        )

    if any(
        termino in texto_base
        for termino in [
            "software",
            "aplicación",
            "plataforma",
            "contenido",
            "diseño",
            "video",
            "imagen",
            "manual",
            "documentación",
        ]
    ):
        normas.extend(
            [
                {
                    "ambito": "Colombia",
                    "norma": "Ley 23 de 1982 y Ley 1915 de 2018",
                    "entidad": "Congreso de Colombia",
                    "aplicacion": (
                        "Orientan la protección de software, documentación, "
                        "contenidos, diseños y demás obras originales desarrolladas."
                    ),
                    "caracter_aplicacion": (
                        "Obligatoria en materia de derecho de autor"
                    ),
                    "fuente_oficial": (
                        "https://www.suin-juriscol.gov.co/"
                        "viewDocument.asp?id=30035790"
                    ),
                },
                {
                    "ambito": "Internacional",
                    "norma": "Decisión Andina 351 de 1993",
                    "entidad": "Comunidad Andina",
                    "aplicacion": (
                        "Establece el régimen común de derecho de autor y derechos "
                        "conexos aplicable a obras del ingenio."
                    ),
                    "caracter_aplicacion": (
                        "Norma supranacional aplicable en Colombia"
                    ),
                    "fuente_oficial": (
                        "https://www.comunidadandina.org/"
                        "StaticFiles/DocOf/DEC351.pdf"
                    ),
                },
            ]
        )

    if any(
        termino in texto_base
        for termino in [
            "venta",
            "cliente",
            "consumidor",
            "comercialización",
            "producto",
            "servicio",
            "garantía",
        ]
    ):
        normas.append(
            {
                "ambito": "Colombia",
                "norma": "Ley 1480 de 2011",
                "entidad": "Congreso de Colombia",
                "aplicacion": (
                    "Debe revisarse cuando el producto o servicio se ofrezca a "
                    "consumidores, especialmente frente a información, seguridad, "
                    "calidad, idoneidad y garantías."
                ),
                "caracter_aplicacion": (
                    "Obligatoria cuando exista una relación de consumo"
                ),
                "fuente_oficial": (
                    "https://www.suin-juriscol.gov.co/"
                    "viewDocument.asp?id=1681955"
                ),
            }
        )

    if any(
        termino in texto_base
        for termino in [
            "mensaje de datos",
            "firma electrónica",
            "comercio electrónico",
            "transacción",
            "pago",
            "documento electrónico",
        ]
    ):
        normas.append(
            {
                "ambito": "Colombia",
                "norma": "Ley 527 de 1999",
                "entidad": "Congreso de Colombia",
                "aplicacion": (
                    "Regula mensajes de datos, comercio electrónico y firmas "
                    "digitales cuando estos elementos formen parte de la solución."
                ),
                "caracter_aplicacion": (
                    "Obligatoria cuando se utilicen mensajes de datos o firmas digitales"
                ),
                "fuente_oficial": (
                    "https://www.suin-juriscol.gov.co/"
                    "viewDocument.asp?id=1662013"
                ),
            }
        )

    if normas and not any(
        item.get("ambito") == "Internacional"
        for item in normas
    ):
        normas.append(
            {
                "ambito": "Internacional",
                "norma": "ISO 9001:2015",
                "entidad": "International Organization for Standardization",
                "aplicacion": (
                    "Puede utilizarse como referencia para organizar procesos, "
                    "controlar la calidad, documentar requisitos y promover la "
                    "mejora continua del producto o servicio."
                ),
                "caracter_aplicacion": (
                    "Referencia técnica voluntaria, salvo exigencia contractual"
                ),
                "fuente_oficial": (
                    "https://www.iso.org/standard/62085.html"
                ),
            }
        )

    if not normas:
        normas = [
            {
                "ambito": "Colombia",
                "norma": "Ley 23 de 1982",
                "entidad": "Congreso de Colombia",
                "aplicacion": (
                    "Debe considerarse para proteger la documentación, planos, "
                    "textos, gráficos y demás creaciones originales del proyecto."
                ),
                "caracter_aplicacion": (
                    "Obligatoria en materia de derecho de autor"
                ),
                "fuente_oficial": (
                    "https://www.suin-juriscol.gov.co/"
                    "viewDocument.asp?id=30035790"
                ),
            },
            {
                "ambito": "Internacional",
                "norma": "Decisión Andina 351 de 1993",
                "entidad": "Comunidad Andina",
                "aplicacion": (
                    "Complementa el marco de protección de las obras y contenidos "
                    "originales generados en el proyecto."
                ),
                "caracter_aplicacion": (
                    "Norma supranacional aplicable en Colombia"
                ),
                "fuente_oficial": (
                    "https://www.comunidadandina.org/"
                    "StaticFiles/DocOf/DEC351.pdf"
                ),
            },
        ]

    normas = normalizar_normas_aplicables(normas)

    texto_normativo = (
        "La normatividad aplicable debe interpretarse de acuerdo con la naturaleza "
        "de la solución, los datos tratados, los usuarios, los componentes técnicos "
        "y la etapa de implementación. Las disposiciones colombianas identificadas "
        "establecen obligaciones que pueden resultar exigibles cuando se configure "
        "su supuesto de aplicación, mientras que los estándares internacionales "
        "funcionan como referentes técnicos o buenas prácticas, salvo que hayan sido "
        "adoptados por una regulación, un contrato o un requisito sectorial. La "
        "selección presentada no constituye una declaración de cumplimiento ni "
        "sustituye la revisión especializada previa a la puesta en operación, "
        "comercialización, certificación o escalamiento del proyecto. En cada fase "
        "de continuidad deben verificarse la vigencia de las disposiciones, la "
        "autoridad competente, los permisos aplicables, la gestión de evidencias y "
        "las responsabilidades de quienes operen la solución."
    )

    return {
        "normatividad_aplicable": texto_normativo,
        "normas_aplicables": normas,
    }


def investigar_normatividad_aplicable(
    datos: dict,
    modelo_openai: str,
) -> dict:
    """
    Identifica mediante búsqueda web la normatividad realmente relacionada con
    el proyecto, usando fuentes oficiales colombianas e internacionales.
    """
    if OpenAI is None:
        raise RuntimeError(
            "La librería openai no está instalada y no es posible verificar "
            "la normatividad aplicable."
        )

    api_key = obtener_api_key()

    if not api_key:
        raise RuntimeError(
            "No se encontró OPENAI_API_KEY para investigar la normatividad."
        )

    client = OpenAI(api_key=api_key)

    prompt = f"""
Actúa como investigador de regulación tecnológica y normalización.

Identifica la normatividad que debe considerarse para el siguiente proyecto de
base tecnológica, tanto en Colombia como en el ámbito internacional.

Proyecto: {datos.get('nombre_proyecto', '')}
Descripción: {datos.get('descripcion_general_proyecto', '')}
Línea tecnológica: {datos.get('linea_tecnologica', '')}
Entregables: {datos.get('entregables_proyecto_base', '')}
Innovación: {datos.get('innovacion_proyecto_base', '')}
Actividades:
{actividades_en_texto(datos.get('actividades_ejecutadas_base', []))}

REGLAS
- Selecciona únicamente normas directamente relacionadas con el proyecto.
- Incluye leyes, decretos, resoluciones o reglamentos colombianos vigentes.
- Incluye normas o estándares internacionales solo cuando sean pertinentes.
- Distingue claramente entre obligación legal, aplicación condicionada,
  norma supranacional y referencia técnica voluntaria.
- No afirmes que el proyecto cumple una norma.
- No inventes números, títulos, entidades, versiones ni enlaces.
- Verifica vigencia y alcance en fuentes oficiales o entidades normalizadoras.
- Para Colombia prioriza SUIN-Juriscol, Función Pública, SIC, MinTIC,
  ministerios, INVIMA y demás autoridades competentes.
- Para el ámbito internacional prioriza Comunidad Andina, ISO, IEC, W3C,
  ITU, OMS, FAO, Codex, OCDE o Naciones Unidas, según el proyecto.
- Entrega entre 4 y 10 normas en total. Si una categoría no aplica, omítela.
- El texto introductorio debe tener entre 260 y 380 palabras y explicar por
  qué el marco seleccionado resulta pertinente, sin convertirlo en asesoría
  jurídica ni repetir toda la lista.

Responde exclusivamente con JSON válido:
{{
  "normatividad_aplicable": "texto técnico de 260 a 380 palabras",
  "normas_aplicables": [
    {{
      "ambito": "Colombia o Internacional",
      "norma": "tipo, número, año y nombre oficial",
      "entidad": "autoridad u organismo responsable",
      "aplicacion": "relación concreta con el proyecto",
      "caracter_aplicacion": "obligatoria, condicionada, supranacional o referencia técnica",
      "fuente_oficial": "https://..."
    }}
  ]
}}
"""

    respuesta = client.responses.create(
        model=modelo_openai,
        tools=[{"type": "web_search"}],
        input=prompt,
        temperature=0.1,
    )

    datos_respuesta = extraer_json_respuesta(
        getattr(respuesta, "output_text", "")
    )

    texto_normativo = limpiar_texto(
        str(datos_respuesta.get("normatividad_aplicable", ""))
    )
    normas = normalizar_normas_aplicables(
        datos_respuesta.get("normas_aplicables", [])
    )

    ambitos = {item["ambito"] for item in normas}

    if (
        not texto_normativo
        or len(normas) < 3
        or "Colombia" not in ambitos
        or "Internacional" not in ambitos
    ):
        raise RuntimeError(
            "La búsqueda no produjo una selección normativa suficiente y "
            "verificable. Intenta nuevamente o amplía la descripción del proyecto."
        )

    return {
        "normatividad_aplicable": texto_normativo,
        "normas_aplicables": normas,
    }


def referencias_bibliograficas_proyecto(datos: dict) -> list[str]:
    referencias = [
        limpiar_texto(item.get("referencia_apa", ""))
        for item in datos.get("referentes_estado_arte", [])
        if limpiar_texto(item.get("referencia_apa", ""))
    ]

    anio = datos.get("fecha_entrega", date.today()).year
    referencias.append(
        (
            f"Red Tecnoparque SENA. ({anio}). Estado del Arte del proyecto "
            f"{datos.get('nombre_proyecto', '')}. Documento de planeación del proyecto."
        )
    )

    normatividad = datos.get("normatividad_investigada", {})
    for item in normalizar_normas_aplicables(
        normatividad.get("normas_aplicables", [])
        if isinstance(normatividad, dict)
        else []
    ):
        referencias.append(
            f"{item['entidad']}. {item['norma']}. {item['fuente_oficial']}"
        )

    return list(dict.fromkeys(referencias))


def anexos_manuales() -> list[str]:
    return [
        "Espacio para insertar manualmente enlaces de evidencias y repositorios.",
        "Espacio para insertar manualmente imágenes, diagramas, planos o capturas.",
        "Espacio para relacionar manualmente otros soportes técnicos del proyecto.",
    ]


def recomendacion_continuidad_trl(trl_alcanzado: str) -> str:
    coincidencia = re.search(r"(\d+)", str(trl_alcanzado or ""))

    if not coincidencia:
        return (
            "Se recomienda definir una fase posterior de validación y continuidad "
            "de acuerdo con el nivel de madurez tecnológica comprobado."
        )

    nivel = int(coincidencia.group(1))

    if 1 <= nivel < 9:
        return (
            f"Al finalizar en TRL {nivel}, es viable formular un nuevo proyecto "
            f"orientado a avanzar hacia TRL {nivel + 1}, sin afirmar que dicho "
            "nivel haya sido alcanzado en la ejecución actual."
        )

    return (
        "Al finalizar en TRL 9, se recomienda concentrar la continuidad en adopción, "
        "escalamiento, sostenibilidad, transferencia y seguimiento del desempeño."
    )


def propiedad_intelectual_modo_prueba(datos: dict) -> str:
    texto_base = " ".join(
        [
            datos.get("descripcion_general_proyecto", ""),
            datos.get("entregables_proyecto_base", ""),
            datos.get("innovacion_proyecto_base", ""),
            actividades_en_texto(
                datos.get("actividades_ejecutadas_base", [])
            ),
        ]
    ).casefold()

    mecanismos: list[str] = []

    if any(
        palabra in texto_base
        for palabra in [
            "software", "aplicación", "plataforma", "código",
            "algoritmo", "base de datos", "sistema web",
        ]
    ):
        mecanismos.append(
            "registro de software ante la Dirección Nacional de Derecho de "
            "Autor, acompañado por la protección del código, la documentación "
            "y los contenidos originales mediante derecho de autor"
        )

    if any(
        palabra in texto_base
        for palabra in [
            "mejora funcional", "mecanismo", "dispositivo", "equipo",
            "prototipo físico", "máquina",
        ]
    ):
        mecanismos.append(
            "modelo de utilidad ante la Superintendencia de Industria y "
            "Comercio, siempre que la configuración funcional cumpla los "
            "requisitos de novedad y aplicación industrial"
        )

    if any(
        palabra in texto_base
        for palabra in [
            "invención", "nuevo principio", "solución técnica nueva",
            "procedimiento novedoso",
        ]
    ):
        mecanismos.append(
            "patente de invención ante la Superintendencia de Industria y "
            "Comercio, condicionada a una búsqueda de antecedentes y al "
            "cumplimiento de novedad, nivel inventivo y aplicación industrial"
        )

    if any(
        palabra in texto_base
        for palabra in [
            "apariencia", "forma externa", "carcasa", "diseño industrial",
            "configuración estética",
        ]
    ):
        mecanismos.append(
            "registro de diseño industrial ante la Superintendencia de "
            "Industria y Comercio para la apariencia externa original"
        )

    if any(
        palabra in texto_base
        for palabra in [
            "marca", "logo", "identidad visual", "nombre comercial",
            "signo distintivo",
        ]
    ):
        mecanismos.append(
            "registro de marca ante la Superintendencia de Industria y "
            "Comercio para los signos distintivos efectivamente utilizados"
        )

    if any(
        palabra in texto_base
        for palabra in [
            "fórmula", "receta", "know-how", "proceso confidencial",
            "secreto", "parámetro reservado",
        ]
    ):
        mecanismos.append(
            "secreto empresarial para la información técnica o comercial que "
            "se mantenga reservada mediante controles y acuerdos de confidencialidad"
        )

    if not mecanismos:
        mecanismos.append(
            "derecho de autor sobre la documentación técnica, planos, textos, "
            "gráficos y demás obras originales desarrolladas"
        )

    mecanismos = mecanismos[:2]

    return (
        "El análisis automático de la naturaleza de la solución, sus entregables, "
        "la innovación reportada y las actividades ejecutadas permite identificar "
        "como mecanismos pertinentes "
        + "; y ".join(mecanismos)
        + ". La selección definitiva requiere verificar la titularidad, la novedad, "
        "el nivel de divulgación y los antecedentes aplicables antes de presentar "
        "una solicitud. Para la transferencia tecnológica se recomienda organizar "
        "los activos, conservar evidencias de autoría, definir condiciones de uso y "
        "establecer acuerdos de licencia, colaboración o adopción con los aliados "
        "que participen en una fase posterior."
    )


def _complemento_apartado(clave: str, datos: dict) -> str:
    metodologia = metodologias_en_texto(datos)
    entregables = datos.get("entregables_proyecto_base", "")
    innovacion = datos.get("innovacion_proyecto_base", "")
    impacto = datos.get("impacto_proyecto_base", "")
    complementos = {
        "introduccion": (
            "El alcance del informe se limita a la experiencia efectivamente "
            "desarrollada y al nivel de madurez tecnológica declarado. La lectura "
            "del documento debe permitir reconocer el propósito, los actores, el "
            "entorno de aplicación y la relación entre la necesidad y la solución, "
            "sin anticipar detalles que corresponden a resultados, viabilidad o "
            "propiedad intelectual. Esta delimitación facilita una presentación "
            "ordenada del cierre y mantiene la trazabilidad de la información."
        ),
        "planteamiento_problema": (
            "La formulación del problema debe diferenciar causas, efectos y usuarios "
            "afectados, además de precisar la brecha que justificó la intervención. "
            "El análisis no convierte la solución en el centro del apartado; se "
            "concentra en las condiciones iniciales, las restricciones observadas y "
            "las consecuencias de mantener la situación sin una respuesta técnica "
            "adaptada al contexto."
        ),
        "estado_arte_tecnica": (
            "La comparación tecnológica considera el propósito, la arquitectura, la "
            "madurez, la accesibilidad y la capacidad de adaptación de las soluciones "
            "existentes. Los referentes se utilizan para identificar tendencias y "
            "diferencias, no para afirmar equivalencias completas. El aporte innovador "
            f"del proyecto se relaciona con {innovacion}, de acuerdo con la información "
            "suministrada y sin atribuir características que no hayan sido verificadas."
        ),
        "metodologia_desarrollo": (
            f"La metodología seleccionada fue {metodologia}. Su aplicación se interpreta "
            "como una secuencia de comprensión, definición, diseño, implementación, "
            "revisión y ajuste. La trazabilidad entre decisiones y actividades permite "
            "explicar cómo se transformaron los requerimientos en componentes y cómo "
            "se controlaron los cambios durante el desarrollo."
        ),
        "desarrollo_proyecto": (
            f"El desarrollo debe explicar cómo las metodologías {metodologia} orientaron "
            "la ejecución del proyecto y cómo sus principios se aplicaron durante las "
            "fases de análisis, planeación, diseño, construcción, integración, revisión "
            "y ajuste. Las actividades se incorporan mediante una narración técnica "
            "continua, previamente corregida y reformulada, sin reproducir literalmente "
            "el texto diligenciado por el usuario ni convertir el apartado en una lista. "
            "Cada acción debe relacionarse con decisiones, procedimientos, componentes, "
            "validaciones o avances concretos dentro de la solución."
        ),
        "resultados_obtenidos": (
            "La valoración de los resultados se realiza por su correspondencia con los "
            "objetivos, su función dentro de la solución y su aporte al nivel TRL "
            "alcanzado. La tabla de entregables deja disponibles espacios de evidencia "
            "para incorporar enlaces en Word sin exigir archivos durante el "
            "diligenciamiento."
        ),
        "analisis_viabilidad": (
            "La viabilidad se examina desde condiciones técnicas, operativas, económicas, "
            "normativas y de adopción. El análisis distingue capacidades disponibles de "
            "requisitos futuros, identifica dependencias y reconoce limitaciones que "
            "deben resolverse antes de un escalamiento. También considera mantenimiento, "
            "documentación, soporte, alianzas y recursos como factores de continuidad."
        ),
        "propiedad_transferencia": (
            "La recomendación de protección no equivale a un derecho concedido. Antes de "
            "tramitar cualquier mecanismo deben revisarse antecedentes, titularidad, "
            "divulgaciones previas y acuerdos entre participantes. La transferencia puede "
            "estructurarse mediante licencias, colaboración, cesión o adopción, según el "
            "activo consolidado y la estrategia de continuidad del proyecto."
        ),
        "impacto_proyecto": (
            f"El impacto reportado fue: {impacto} Su análisis debe identificar quiénes "
            "reciben el beneficio, mediante qué cambio concreto y bajo cuáles condiciones "
            "puede sostenerse. Se diferencian efectos observados de expectativas futuras y "
            "se evita presentar cifras o resultados que no hayan sido suministrados. La "
            "continuidad del impacto depende de adopción, soporte y seguimiento."
        ),
        "conclusiones": (
            "Las conclusiones articulan pertinencia, aprendizaje, entregables, innovación "
            "y limitaciones sin introducir hechos nuevos. El nivel TRL se utiliza como "
            "referencia para orientar una fase posterior, manteniendo explícito que el "
            "siguiente nivel constituye una meta de continuidad y no un resultado ya "
            "alcanzado dentro del proyecto que se cierra."
        ),
    }

    return complementos.get(clave, "")


def _ampliacion_adicional(clave: str, datos: dict) -> str:
    ampliaciones = {
        "introduccion": (
            "El cierre documental también permite ubicar el proyecto dentro del proceso "
            "de acompañamiento de TecnoParque, reconocer la participación del talento y "
            "del experto, y establecer el periodo al que corresponden los resultados. "
            "La introducción delimita el objeto del informe y aclara que las valoraciones "
            "posteriores se basan en la información efectivamente suministrada. De esta "
            "forma, el lector obtiene una visión inicial suficiente para comprender el "
            "sentido del proyecto, su relación con el entorno y la razón por la cual se "
            "documenta el proceso de desarrollo. La sección evita convertir el contexto "
            "en una repetición de actividades o productos, y reserva para cada numeral "
            "los elementos que le corresponden. Esta organización fortalece la coherencia "
            "del documento y facilita la revisión de la trazabilidad entre necesidad, "
            "propósito, ejecución, resultados e impacto."
        ),
        "planteamiento_problema": (
            "La situación inicial debe comprenderse a partir de las prácticas existentes, "
            "las limitaciones de acceso, desempeño, integración o disponibilidad y las "
            "dificultades que enfrentaban los usuarios. El apartado identifica por qué "
            "esas condiciones afectaban el proceso y qué consecuencias podían mantenerse "
            "si no se intervenían. También diferencia síntomas de causas para evitar una "
            "formulación superficial. La justificación tecnológica surge cuando la brecha "
            "no puede resolverse únicamente con ajustes menores y requiere diseño, "
            "integración o validación de una alternativa. La redacción conserva una relación "
            "directa con el contexto reportado, pero no atribuye magnitudes, frecuencias ni "
            "pérdidas que no hayan sido documentadas. Así, el problema funciona como base "
            "para valorar la pertinencia de los objetivos y la correspondencia de los "
            "resultados obtenidos."
        ),
        "estado_arte_tecnica": (
            "El análisis de referentes debe valorar el grado de madurez de las alternativas, "
            "las tecnologías empleadas, el tipo de usuario atendido y las condiciones de "
            "operación. La comparación no se limita a enumerar productos, sino que identifica "
            "tendencias de diseño, integración, automatización, interoperabilidad, usabilidad "
            "y sostenibilidad. Los dos casos reales se utilizan como evidencia de avances "
            "existentes y como punto de contraste para reconocer el aporte diferencial de "
            "la solución desarrollada. También se consideran limitaciones de adaptación, "
            "costos de adopción, dependencia de infraestructura y posibilidades de mejora, "
            "siempre que puedan deducirse razonablemente de las fuentes consultadas. Esta "
            "síntesis orienta la lectura técnica del informe, mientras la revisión extensa, "
            "las matrices comparativas y las demás referencias permanecen en el documento "
            "de planeación correspondiente."
        ),
        "metodologia_desarrollo": (
            "La metodología se evidencia en la manera de organizar decisiones, priorizar "
            "requerimientos y transformar observaciones en ajustes concretos. El proceso "
            "puede integrar exploración con usuarios, definición de criterios, construcción "
            "progresiva, integración de componentes y revisión del funcionamiento. Cada "
            "etapa conserva relación con una actividad reportada y con un resultado esperado, "
            "lo que permite explicar la evolución sin inventar ceremonias, roles o documentos "
            "que no hayan existido. La iteración se interpreta como un mecanismo para reducir "
            "incertidumbre, detectar incompatibilidades y mejorar la solución antes del cierre. "
            "La validación se presenta de acuerdo con las evidencias disponibles y el TRL, "
            "diferenciando pruebas realizadas de evaluaciones que deberán completarse en una "
            "fase posterior. Así, el enfoque metodológico ofrece trazabilidad y no una simple "
            "lista de conceptos."
        ),
        "desarrollo_proyecto": (
            "La narración del desarrollo explica la secuencia temporal y técnica de las "
            "acciones, señalando cómo los insumos de una actividad permitieron ejecutar la "
            "siguiente. Se describen los momentos de análisis, diseño, preparación, construcción, "
            "configuración, integración y revisión que resulten aplicables. Las decisiones "
            "se presentan por su efecto sobre el funcionamiento, la compatibilidad o la "
            "adaptación al contexto, sin atribuir especificaciones no reportadas. Cuando se "
            "produjeron ajustes, se explica su relación con hallazgos del proceso y con la "
            "necesidad de mejorar el prototipo. La sección diferencia claramente las acciones "
            "ejecutadas de los productos obtenidos y del impacto posterior. Esta estructura "
            "facilita reconocer la contribución de cada actividad y deja una base coherente "
            "para asociar evidencias en la tabla de resultados."
        ),
        "resultados_obtenidos": (
            "Los resultados se describen como productos verificables del proceso y no como "
            "expectativas. El análisis establece qué función cumple cada entregable, cómo se "
            "relaciona con los objetivos y qué parte de la necesidad permite atender. También "
            "reconoce el aporte de la innovación cuando esta se manifiesta en una mejora, una "
            "integración particular o una adaptación al entorno. La valoración respeta el TRL "
            "alcanzado y evita presentar como definitivo aquello que todavía requiere pruebas "
            "en condiciones más amplias. La tabla complementaria organiza las actividades y "
            "deja un espacio para enlaces de evidencia, de modo que fotografías, repositorios, "
            "actas, manuales o demostraciones puedan incorporarse posteriormente. Esta separación "
            "entre síntesis y soportes mantiene el informe legible y facilita su verificación."
        ),
        "analisis_viabilidad": (
            "La viabilidad técnica considera disponibilidad, compatibilidad, mantenibilidad y "
            "posibilidad de reproducir o escalar la solución. La dimensión operativa examina "
            "responsables, capacidades, rutinas y condiciones de uso. El componente económico "
            "identifica recursos que deberán estimarse sin inventar valores, mientras el análisis "
            "normativo reconoce obligaciones aplicables a datos, seguridad, propiedad intelectual, "
            "salud, ambiente o calidad, según la naturaleza del proyecto. La adopción depende de "
            "que los usuarios comprendan el funcionamiento y perciban beneficios suficientes. "
            "También se valoran sostenibilidad, soporte, disponibilidad de insumos, alianzas y "
            "oportunidades de mercado cuando sean pertinentes. El apartado concluye diferenciando "
            "condiciones favorables, limitaciones y acciones necesarias antes de una implementación "
            "más amplia."
        ),
        "propiedad_transferencia": (
            "El mecanismo sugerido debe corresponder a la naturaleza del activo y no a una lista "
            "general de posibilidades. Los componentes originales se analizan para distinguir "
            "obras protegibles por derecho de autor, desarrollos funcionales con potencial de "
            "propiedad industrial, signos distintivos y conocimiento reservado. La documentación "
            "de fechas, autores, versiones y aportes resulta esencial para sustentar titularidad. "
            "También deben evitarse divulgaciones que puedan afectar la novedad cuando se estudie "
            "una patente, un modelo de utilidad o un diseño industrial. La estrategia de transferencia "
            "puede involucrar licenciamiento, colaboración, adopción institucional, prestación de "
            "servicios o continuidad conjunta, siempre con reglas claras sobre uso, mantenimiento, "
            "confidencialidad y distribución de beneficios. La recomendación final permanece sujeta "
            "a una revisión jurídica y técnica específica."
        ),
        "impacto_proyecto": (
            "El impacto se interpreta a partir del cambio que la solución puede producir en usuarios, "
            "procesos u organizaciones. La dimensión tecnológica se relaciona con capacidades, acceso, "
            "integración o apropiación; la productiva, con mejoras en operación o toma de decisiones; "
            "la social, con beneficios para personas o comunidades; la económica, con oportunidades de "
            "eficiencia o generación de valor; y la ambiental, con reducción o mejor gestión de recursos, "
            "cuando aplique. La redacción distingue beneficios observados de efectos potenciales y evita "
            "cuantificaciones no respaldadas. También identifica condiciones para sostener el impacto, "
            "como capacitación, soporte, mantenimiento, actualización, participación de beneficiarios y "
            "seguimiento. Esta perspectiva permite valorar el aporte sin confundirlo con la simple entrega "
            "de un prototipo."
        ),
        "conclusiones": (
            "El cierre reconoce qué aspectos quedaron consolidados, cuáles requieren validación adicional "
            "y qué aprendizajes deben conservarse. La conclusión no repite el desarrollo completo; relaciona "
            "la necesidad, los objetivos, los entregables y el impacto para valorar la coherencia global. "
            "También identifica limitaciones técnicas, operativas o documentales que pueden orientar una "
            "fase posterior. La recomendación de continuidad se formula de acuerdo con el TRL alcanzado y "
            "presenta el siguiente nivel como una meta condicionada a nuevas actividades, recursos y evidencias. "
            "Se consideran además oportunidades de protección, transferencia, adopción o escalamiento, sin "
            "afirmar que ya se hayan materializado. De esta manera, el informe finaliza con una valoración "
            "realista del estado actual y con una orientación concreta para la evolución responsable del proyecto."
        ),
    }
    return ampliaciones.get(clave, "")


def ajustar_rango_palabras(
    texto: str,
    clave: str,
    datos: dict,
    minimo: int = 300,
    maximo: int = 340,
) -> str:
    resultado = limpiar_texto(texto)

    if clave == "desarrollo_proyecto":
        minimo = 120
        maximo = 190

    for fragmento in [
        _complemento_apartado(clave, datos),
        _ampliacion_adicional(clave, datos),
    ]:
        if len(resultado.split()) >= minimo:
            break
        fragmento_limpio = limpiar_texto(fragmento)
        if fragmento_limpio and fragmento_limpio.casefold() not in resultado.casefold():
            resultado = f"{resultado} {fragmento_limpio}".strip()

    if len(resultado.split()) < minimo:
        cierres_finales = {
            "introduccion": (
                "Con esta delimitación, el apartado ofrece una entrada suficiente para comprender "
                "el documento, ubicar el proyecto en su contexto institucional y reconocer el alcance "
                "del cierre sin confundirlo con una propuesta futura ni con una descripción exhaustiva "
                "de los productos obtenidos."
            ),
            "planteamiento_problema": (
                "Esta formulación permite usar el problema como criterio de evaluación del proyecto: "
                "los objetivos deben responder a la brecha identificada y los resultados deben mostrar "
                "una contribución verificable a su atención, aun cuando permanezcan necesidades de "
                "validación o ampliación."
            ),
            "estado_arte_tecnica": (
                "La síntesis resultante ofrece un marco suficiente para interpretar la posición del "
                "proyecto frente a soluciones existentes y para comprender por qué su adaptación al "
                "entorno puede representar una mejora pertinente, sin sustituir la revisión académica "
                "y técnica completa realizada en planeación."
            ),
            "metodologia_desarrollo": (
                "La explicación metodológica conserva correspondencia con lo ejecutado y evita presentar "
                "el proyecto como si hubiera aplicado procedimientos no reportados. Su valor está en mostrar "
                "cómo se ordenó el trabajo, cómo se revisaron los avances y cómo las decisiones permitieron "
                "consolidar los entregables."
            ),
            "desarrollo_proyecto": (
                "El resultado es una descripción continua del proceso, útil para que terceros comprendan "
                "la evolución del proyecto, identifiquen los momentos principales y relacionen cada actividad "
                "con la solución final sin perder claridad entre ejecución, resultados e impacto."
            ),
            "resultados_obtenidos": (
                "La valoración final mantiene una relación directa con los productos reportados y deja claro "
                "qué fue alcanzado dentro del proyecto. Cualquier desempeño adicional, adopción masiva o "
                "escalamiento deberá demostrarse en fases posteriores mediante evidencias específicas."
            ),
            "analisis_viabilidad": (
                "Con base en estas dimensiones, la continuidad puede planearse mediante una evaluación gradual "
                "de recursos, riesgos y condiciones de adopción. La decisión de implementar o escalar debe "
                "apoyarse en pruebas adicionales y en una estimación formal de costos, responsabilidades y "
                "requisitos aplicables."
            ),
            "propiedad_transferencia": (
                "Esta orientación permite priorizar acciones realistas de protección y transferencia, reducir "
                "riesgos de pérdida de novedad o autoría y preparar una estrategia compatible con la etapa de "
                "madurez del proyecto y con los intereses de los participantes."
            ),
            "impacto_proyecto": (
                "La valoración del impacto debe actualizarse cuando existan nuevas evidencias de uso, cobertura "
                "o desempeño. Por ahora, el informe conserva los efectos reportados y presenta las condiciones "
                "que permitirían mantener o ampliar los beneficios en una fase posterior."
            ),
            "conclusiones": (
                "Esta lectura de cierre permite tomar decisiones de continuidad sin sobreestimar el resultado. "
                "El proyecto queda documentado en su nivel real de avance y con una ruta de trabajo que puede "
                "convertirse en una nueva formulación orientada a validaciones más exigentes."
            ),
        }
        resultado = f"{resultado} {cierres_finales.get(clave, '')}".strip()

    if len(resultado.split()) < minimo:
        refuerzos_finales = {
            "introduccion": (
                "La información institucional, el TRL y la fecha de entrega completan esta "
                "ubicación inicial y permiten interpretar el documento como un cierre correspondiente "
                "a una etapa concreta del proceso tecnológico. El apartado queda así concentrado en "
                "presentar el proyecto y orientar la lectura de los numerales siguientes."
            ),
            "planteamiento_problema": (
                "La delimitación obtenida evita formulaciones demasiado amplias y permite reconocer "
                "una necesidad susceptible de ser atendida mediante un proyecto de base tecnológica. "
                "Esta precisión es indispensable para comprobar posteriormente la relación entre el "
                "problema, los objetivos formulados y los productos efectivamente alcanzados."
            ),
            "estado_arte_tecnica": (
                "La revisión también ayuda a reconocer oportunidades de interoperabilidad, accesibilidad, "
                "mantenimiento y escalamiento que pueden orientar nuevas decisiones. La comparación se "
                "mantiene sustentada en fuentes verificables y evita convertir tendencias generales en "
                "resultados propios del proyecto."
            ),
            "metodologia_desarrollo": (
                "La secuencia metodológica facilita además documentar responsables, insumos, criterios de "
                "revisión y decisiones de ajuste. Esta trazabilidad resulta útil para reproducir el proceso, "
                "identificar aprendizajes y preparar una eventual fase de continuidad con actividades y "
                "evidencias mejor delimitadas."
            ),
            "desarrollo_proyecto": (
                "La articulación entre actividades evidencia una progresión desde la definición hasta la "
                "consolidación de los productos. El relato técnico conserva el orden del trabajo y permite "
                "comprender qué decisiones fueron necesarias para integrar componentes, resolver dificultades "
                "y preparar la solución para su valoración final."
            ),
            "resultados_obtenidos": (
                "La síntesis permite reconocer el alcance real de los productos sin confundirlos con impactos "
                "futuros. Los espacios de evidencia facilitan completar la trazabilidad documental después de "
                "la generación del archivo, conservando flexibilidad para incorporar enlaces, repositorios, "
                "actas, fotografías o registros de prueba."
            ),
            "analisis_viabilidad": (
                "La conclusión de viabilidad debe entenderse como una valoración técnica preliminar basada en "
                "el estado actual del proyecto. Una decisión de adopción o inversión requerirá complementar "
                "esta información con pruebas de desempeño, costos detallados, revisión regulatoria, análisis "
                "de usuarios y definición de responsabilidades operativas. También será necesario establecer "
                "un plan de mantenimiento, soporte, actualización y gestión de riesgos que permita sostener "
                "el funcionamiento en condiciones reales durante una etapa posterior."
            ),
            "propiedad_transferencia": (
                "La identificación temprana de los activos facilita organizar expedientes, acuerdos y evidencias "
                "antes de una negociación o divulgación. También permite seleccionar aliados y mecanismos de "
                "transferencia compatibles con la estrategia de desarrollo, evitando comprometer derechos o "
                "información reservada sin una evaluación previa."
            ),
            "impacto_proyecto": (
                "La lectura de impacto conserva una relación directa con los beneficiarios y con el problema "
                "atendido. Su seguimiento posterior permitirá confirmar permanencia, ampliar la cobertura y "
                "establecer indicadores, siempre que se definan mecanismos de recolección de evidencias y una "
                "responsabilidad clara sobre la continuidad de la solución."
            ),
            "conclusiones": (
                "El informe deja así una base para decidir si conviene fortalecer el prototipo, ampliar las "
                "pruebas, preparar una estrategia de adopción o formular un nuevo proyecto. La recomendación "
                "debe revisarse junto con los recursos disponibles, los aliados y las condiciones necesarias "
                "para demostrar el siguiente nivel de madurez."
            ),
        }
        resultado = f"{resultado} {refuerzos_finales.get(clave, '')}".strip()

    palabras = resultado.split()
    if len(palabras) > maximo:
        resultado = " ".join(palabras[:maximo]).rstrip(" ,;:")
        if resultado[-1:] not in ".!?":
            resultado += "."

    return resultado


def asegurar_nota_estado_arte(texto: str) -> str:
    nota = (
        "El Estado del Arte completo se encuentra en la sección de documentos "
        "de planeación del proyecto, específicamente en el documento Estado del Arte."
    )

    if "documento estado del arte" not in texto.casefold():
        return f"{texto.strip()}\n\n{nota}"

    return texto.strip()


def contenido_modo_prueba(datos: dict) -> dict:
    metodologia = metodologias_en_texto(datos)
    actividades_corregidas = [
        corregir_actividad_basica(actividad)
        for actividad in datos.get("actividades_ejecutadas_base", [])
        if corregir_actividad_basica(actividad)
    ]
    entregables_corregidos = [
        corregir_entregable_basico(entregable)
        for entregable in datos.get("entregables_proyecto_lista", [])
        if corregir_entregable_basico(entregable)
    ]
    referentes = datos.get("referentes_estado_arte", [])

    referentes_texto = " ".join(
        (
            f"{item.get('nombre', '')}, desarrollado por "
            f"{item.get('entidad', '')}, constituye un referente porque "
            f"{item.get('descripcion', '')} {item.get('cita_corta', '')}."
        )
        for item in referentes
    )

    descripcion = datos.get("descripcion_general_proyecto", "")
    entregables = datos.get("entregables_proyecto_base", "")
    innovacion = datos.get("innovacion_proyecto_base", "")
    impacto = datos.get("impacto_proyecto_base", "")

    base = {
        "introduccion": (
            f"El presente informe documenta el cierre técnico del proyecto "
            f"{datos.get('nombre_proyecto', '')}, desarrollado en "
            f"{datos.get('tecnoparque', '')}. La iniciativa se contextualiza a "
            f"partir de la siguiente información: {descripcion} El apartado presenta "
            "el propósito general, los actores, el entorno y el alcance alcanzado, "
            "sin desarrollar de manera anticipada el problema, la metodología ni los "
            "resultados. La organización del informe permite conservar trazabilidad "
            "entre la necesidad, los objetivos, las actividades y el nivel de madurez "
            "tecnológica declarado."
        ),
        "planteamiento_problema": (
            f"La problemática se deriva de las condiciones descritas para el proyecto: "
            f"{descripcion} El análisis identifica la brecha existente, los usuarios o "
            "procesos afectados, las causas que limitaron una respuesta adecuada y las "
            "consecuencias de mantener la situación inicial. La justificación se centra "
            "en la necesidad de disponer de una alternativa tecnológica pertinente y "
            "adaptada al contexto, sin convertir este apartado en una descripción de la "
            "solución o de los productos finales."
        ),
        "objetivo_general": (
            "Desarrollar y validar una solución de base tecnológica que responda a la "
            "necesidad identificada mediante actividades estructuradas de diseño, "
            "implementación, evaluación y documentación."
        ),
        "objetivos_especificos": [
            "Caracterizar la necesidad, los usuarios, los requerimientos y las condiciones de aplicación de la solución.",
            "Diseñar los componentes y criterios técnicos necesarios para materializar la propuesta tecnológica.",
            "Implementar e integrar la solución mediante actividades organizadas de construcción, configuración y ajuste.",
            "Validar los entregables obtenidos y establecer oportunidades de mejora, continuidad y escalamiento.",
        ],
        "estado_arte_tecnica": asegurar_nota_estado_arte(
            "Los avances tecnológicos del área muestran una tendencia hacia soluciones "
            "más modulares, interoperables, accesibles y orientadas al usuario. "
            + referentes_texto
            + f" El aporte innovador informado para el proyecto corresponde a {innovacion}. "
            "La comparación permite identificar mejoras de adaptación, integración o "
            "aplicación frente a alternativas existentes, sin afirmar equivalencias ni "
            "atribuir ventajas que no hayan sido verificadas."
        ),
        "metodologia_desarrollo": metodologia_modo_prueba(datos),
        "actividades_corregidas": actividades_corregidas,
        "entregables_corregidos": entregables_corregidos,
        "desarrollo_proyecto": (
            f"El desarrollo del proyecto se organizó mediante "
            f"{metodologias_en_texto(datos)}, articulando sus fases con las acciones "
            "ejecutadas durante el proceso. Cada actividad se presenta de manera "
            "independiente, con su relación metodológica, descripción técnica y espacio "
            "destinado a incorporar las evidencias correspondientes. Esta estructura "
            "permite documentar la secuencia real del trabajo y facilitar la inclusión "
            "posterior de fotografías, capturas de pantalla, enlaces y demás soportes."
        ),
        "desarrollo_actividades": desarrollo_actividades_modo_prueba(
            datos,
            actividades_corregidas,
        ),
        "normatividad_aplicable": (
            datos.get("normatividad_investigada", {})
            or normatividad_modo_prueba(datos)
        ).get("normatividad_aplicable", ""),
        "normas_aplicables": (
            datos.get("normatividad_investigada", {})
            or normatividad_modo_prueba(datos)
        ).get("normas_aplicables", []),
        "resultados_obtenidos": resultados_modo_prueba_sin_copia(
            datos,
            len(entregables_corregidos),
        ),
        "analisis_viabilidad": (
            "La viabilidad se examina considerando la estabilidad técnica, la operación, "
            "los recursos necesarios, la adopción, el mantenimiento, las obligaciones "
            "normativas y las oportunidades de sostenibilidad. El análisis diferencia "
            "las capacidades comprobadas de los requisitos que deben abordarse en una "
            "fase posterior y evita inventar costos, permisos, ventas o comportamientos "
            "de mercado no suministrados."
        ),
        "propiedad_transferencia": propiedad_intelectual_modo_prueba(datos),
        "impacto_proyecto": (
            f"El impacto informado fue: {impacto} El apartado organiza los beneficios "
            "según las dimensiones tecnológicas, sociales, económicas, ambientales o "
            "productivas que realmente resulten aplicables. Se identifica a los "
            "beneficiarios, el cambio generado y las condiciones necesarias para "
            "sostenerlo, sin atribuir cifras o efectos no reportados."
        ),
        "conclusiones": (
            "El cierre integra la pertinencia de la solución, la relación entre objetivos, "
            "actividades y entregables, el aporte innovador, los aprendizajes y las "
            "limitaciones identificadas. "
            + recomendacion_continuidad_trl(
                datos.get("trl_alcanzado", "")
            )
        ),
        "referencias_bibliograficas": referencias_bibliograficas_proyecto(datos),
        "anexos": anexos_manuales(),
    }

    for clave in [
        "introduccion",
        "planteamiento_problema",
        "estado_arte_tecnica",
        "metodologia_desarrollo",
        "desarrollo_proyecto",
        "resultados_obtenidos",
        "analisis_viabilidad",
        "propiedad_transferencia",
        "impacto_proyecto",
        "conclusiones",
    ]:
        base[clave] = ajustar_rango_palabras(
            base[clave],
            clave,
            datos,
        )

    return base


def normalizar_contenido(
    contenido: object,
    respaldo: dict,
    datos: dict,
) -> dict:
    if not isinstance(contenido, dict):
        contenido = {}

    resultado: dict = {}

    for clave in CLAVES_CONTENIDO:
        valor = contenido.get(clave)

        if clave == "desarrollo_actividades":
            actividades_base = resultado.get(
                "actividades_corregidas",
                respaldo.get("actividades_corregidas", []),
            )
            resultado[clave] = normalizar_desarrollo_actividades(
                valor,
                actividades_base,
                datos,
            )
            continue

        if clave == "normas_aplicables":
            normas = normalizar_normas_aplicables(valor)
            if not normas:
                normas = normalizar_normas_aplicables(
                    respaldo.get("normas_aplicables", [])
                )
            resultado[clave] = normas
            continue

        if clave in {
            "objetivos_especificos",
            "actividades_corregidas",
            "entregables_corregidos",
            "referencias_bibliograficas",
            "anexos",
        }:
            if isinstance(valor, str):
                valor = dividir_lineas(valor)

            if not isinstance(valor, list):
                valor = []

            valor = [
                limpiar_texto(str(item))
                for item in valor
                if limpiar_texto(str(item))
            ]

            if not valor:
                valor = respaldo[clave]

            resultado[clave] = valor
        else:
            texto_valor = limpiar_texto(str(valor or ""))
            if clave in {"metodologia_desarrollo", "desarrollo_proyecto"}:
                texto_valor = limpiar_frases_meta(texto_valor)
            resultado[clave] = texto_valor or respaldo[clave]

    resultado["objetivos_especificos"] = (
        resultado["objetivos_especificos"]
        + respaldo["objetivos_especificos"]
    )[:4]

    cantidad_actividades = len(
        datos.get("actividades_ejecutadas_base", [])
    )
    resultado["actividades_corregidas"] = resultado[
        "actividades_corregidas"
    ][:cantidad_actividades]

    if len(resultado["actividades_corregidas"]) < cantidad_actividades:
        faltantes = datos.get("actividades_ejecutadas_base", [])[
            len(resultado["actividades_corregidas"]):
        ]
        resultado["actividades_corregidas"].extend(
            corregir_actividad_basica(item)
            for item in faltantes
        )

    cantidad_entregables = len(
        datos.get("entregables_proyecto_lista", [])
    )
    resultado["entregables_corregidos"] = resultado[
        "entregables_corregidos"
    ][:cantidad_entregables]

    if len(resultado["entregables_corregidos"]) < cantidad_entregables:
        faltantes_entregables = datos.get(
            "entregables_proyecto_lista",
            [],
        )[len(resultado["entregables_corregidos"]):]

        resultado["entregables_corregidos"].extend(
            corregir_entregable_basico(item)
            for item in faltantes_entregables
        )

    resultado["referencias_bibliograficas"] = (
        referencias_bibliograficas_proyecto(datos)
    )
    resultado["anexos"] = anexos_manuales()

    for clave in [
        "introduccion",
        "planteamiento_problema",
        "estado_arte_tecnica",
        "metodologia_desarrollo",
        "desarrollo_proyecto",
        "resultados_obtenidos",
        "analisis_viabilidad",
        "propiedad_transferencia",
        "impacto_proyecto",
        "conclusiones",
    ]:
        resultado[clave] = ajustar_rango_palabras(
            resultado[clave],
            clave,
            datos,
        )

    resultado["estado_arte_tecnica"] = asegurar_nota_estado_arte(
        resultado["estado_arte_tecnica"]
    )

    continuidad = recomendacion_continuidad_trl(
        datos.get("trl_alcanzado", "")
    )
    if continuidad.casefold() not in resultado["conclusiones"].casefold():
        resultado["conclusiones"] = (
            f"{resultado['conclusiones'].rstrip()} {continuidad}"
        )

    return resultado


def reescribir_desarrollo_sin_copia_literal(
    texto_desarrollo: str,
    actividades_corregidas: list[str],
    actividades_originales: list[str],
    datos: dict,
    modelo_openai: str,
) -> str:
    """
    Reescribe Desarrollo del proyecto cuando detecta fragmentos literales de
    las actividades originales o corregidas.
    """
    metodologia = metodologias_en_texto(datos)

    try:
        respuesta = generar_json_openai(
            instrucciones=(
                "Actúa como redactor técnico senior de TecnoParque SENA. "
                "Reescribe exclusivamente el apartado Desarrollo del proyecto "
                "en español formal, con 450 a 520 palabras. Integra todas las "
                "metodologías indicadas y las actividades ejecutadas dentro de "
                "una narración técnica continua. Corrige ortografía, gramática, "
                "puntuación, coherencia y precisión técnica. No copies, enumeres "
                "ni reproduzcas literalmente las actividades originales o "
                "corregidas. No inventes tecnologías, resultados, cantidades ni "
                "validaciones. Responde exclusivamente en JSON válido con la "
                "clave desarrollo_proyecto."
            ),
            entrada=(
                "METODOLOGÍAS UTILIZADAS\n"
                f"{metodologia}\n\n"
                "ACTIVIDADES CORREGIDAS COMO FUENTE CONCEPTUAL\n"
                f"{actividades_en_texto(actividades_corregidas)}\n\n"
                "TEXTO QUE DEBE REESCRIBIRSE\n"
                f"{texto_desarrollo}\n\n"
                "ESTRUCTURA JSON\n"
                '{"desarrollo_proyecto": "450 a 520 palabras"}'
            ),
            modelo=modelo_openai,
            temperature=0.1,
        )

        if isinstance(respuesta, dict):
            reescrito = limpiar_texto(
                str(respuesta.get("desarrollo_proyecto", ""))
            )
            fuentes = actividades_originales + actividades_corregidas

            if (
                len(reescrito.split()) >= 450
                and not contiene_copia_textual(reescrito, fuentes)
            ):
                return ajustar_rango_palabras(
                    reescrito,
                    "desarrollo_proyecto",
                    datos,
                )

    except Exception:
        pass

    respaldo = contenido_modo_prueba(datos)["desarrollo_proyecto"]
    return ajustar_rango_palabras(
        respaldo,
        "desarrollo_proyecto",
        datos,
    )


def reescribir_resultados_sin_copia_literal(
    texto_resultados: str,
    entregables_corregidos: list[str],
    entregables_originales: list[str],
    datos: dict,
    modelo_openai: str,
) -> str:
    """
    Reescribe Resultados obtenidos cuando detecta coincidencias literales con
    los entregables originales.
    """
    respaldo = resultados_modo_prueba_sin_copia(
        datos,
        len(entregables_corregidos),
    )

    try:
        respuesta = generar_json_openai(
            instrucciones=(
                "Actúa como redactor técnico senior de TecnoParque SENA. "
                "Reescribe exclusivamente el apartado Resultados obtenidos en "
                "español formal, con 300 a 340 palabras. Corrige ortografía, "
                "gramática, puntuación, coherencia y precisión técnica. No "
                "copies literalmente, no enumeres y no reproduzcas frases de "
                "los entregables originales. Analiza su relación con objetivos, "
                "actividades, innovación y TRL. No inventes datos. Responde "
                "exclusivamente en JSON válido con la clave resultados_obtenidos."
            ),
            entrada=(
                "TEXTO QUE DEBE REESCRIBIRSE\n"
                f"{texto_resultados}\n\n"
                "ENTREGABLES CORREGIDOS COMO FUENTE CONCEPTUAL\n"
                f"{actividades_en_texto(entregables_corregidos)}\n\n"
                "TRL ALCANZADO\n"
                f"{datos.get('trl_alcanzado', '')}\n\n"
                "ESTRUCTURA JSON\n"
                '{"resultados_obtenidos": "300 a 340 palabras"}'
            ),
            modelo=modelo_openai,
            temperature=0.1,
        )

        if isinstance(respuesta, dict):
            reescrito = limpiar_texto(
                str(respuesta.get("resultados_obtenidos", ""))
            )

            if (
                reescrito
                and not contiene_copia_textual(
                    reescrito,
                    entregables_originales,
                )
            ):
                return reescrito

    except Exception:
        pass

    return respaldo


def generar_contenido_con_ia(
    datos: dict,
    modelo_openai: str,
) -> dict:
    respaldo = contenido_modo_prueba(datos)
    metodologia = metodologias_en_texto(datos)
    actividades_originales = datos.get(
        "actividades_ejecutadas_base",
        [],
    )
    entregables_originales = datos.get(
        "entregables_proyecto_lista",
        [],
    )
    referentes = datos.get("referentes_estado_arte", [])

    reglas_comunes = """
Actúa como redactor técnico senior de proyectos de base tecnológica de la Red
TecnoParque SENA. Redacta un Informe Final institucional en español formal,
coherente, preciso y verificable.

REGLAS OBLIGATORIAS
- Utiliza únicamente los datos suministrados y los referentes web verificados.
- No inventes cifras, costos, pruebas, certificaciones, normas, patentes,
  registros, clientes, ventas, resultados ni referencias.
- Evita repetir la descripción general del proyecto entre apartados.
- Cada sección debe cumplir únicamente el propósito definido en el formato.
- Cada apartado narrativo debe contener entre 300 y 340 palabras.
- No uses introducciones genéricas ni frases de relleno.
- Corrige ortografía, gramática, puntuación, concordancia y redacción técnica de
  todos los textos suministrados por el usuario antes de incorporarlos.
- No copies literalmente oraciones ni fragmentos extensos de los campos
  diligenciados. Reformula siempre el contenido con lenguaje técnico,
  institucional y claro, conservando estrictamente su significado.
- Las listas destinadas a tablas también deben pasar por corrección y mejora
  técnica, aunque puedan conservar términos propios, nombres o denominaciones.
- No uses markdown.
- Responde exclusivamente en JSON válido.
"""

    contexto = f"""
DATOS GENERALES
Talento: {datos.get('nombre_talento', '')}
Proyecto: {datos.get('nombre_proyecto', '')}
Código: {datos.get('codigo_proyecto', '')}
Experto: {datos.get('nombre_experto', '')}
Línea tecnológica: {datos.get('linea_tecnologica', '')}
TRL inicial: {datos.get('trl_inicial', '')}
TRL alcanzado: {datos.get('trl_alcanzado', '')}
TecnoParque: {datos.get('tecnoparque', '')}

DESCRIPCIÓN GENERAL DEL PROYECTO
{datos.get('descripcion_general_proyecto', '')}

ENTREGABLES OBTENIDOS
{actividades_en_texto(entregables_originales)}

INNOVACIÓN DEL PROYECTO
{datos.get('innovacion_proyecto_base', '')}

ACTIVIDADES EJECUTADAS
{actividades_en_texto(actividades_originales)}

IMPACTO DEL PROYECTO
{datos.get('impacto_proyecto_base', '')}

METODOLOGÍAS SELECCIONADAS
{metodologia}

FASES DE REFERENCIA DE LAS METODOLOGÍAS
{fases_metodologias_en_texto(datos)}

DOS REFERENTES REALES VERIFICADOS PARA EL ESTADO DEL ARTE
{json.dumps(referentes, ensure_ascii=False, indent=2)}
"""

    instrucciones_bloque_1 = reglas_comunes + """
Genera los apartados 2, 3, 4, 5 y 6; corrige las actividades y mejora técnicamente los entregables.

REQUISITOS ESPECÍFICOS
- Introducción: contexto, propósito, entorno, actores y alcance. No desarrolles
  el problema, la metodología, los resultados ni el impacto.
- Planteamiento del problema: necesidad, causas, consecuencias, usuarios
  afectados y justificación. No describas extensamente la solución.
- Objetivo general: una sola oración con verbo en infinitivo.
- Objetivos específicos: exactamente cuatro, cada uno con verbo en infinitivo.
- Estado del arte: resume avances del área e integra exactamente los dos
  referentes reales con sus citas cortas. Explica el aporte innovador y finaliza
  indicando que el Estado del Arte completo está en los documentos de
  planeación, específicamente en el documento Estado del Arte.
- Metodología: utiliza exactamente las metodologías seleccionadas por el usuario.
  Describe técnicamente cada metodología, sus fases, etapas, principios, artefactos o
  ciclos, según corresponda. Para Design Thinking desarrolla empatizar, definir,
  idear, prototipar y evaluar; para modelo en cascada desarrolla requisitos, diseño,
  implementación, pruebas, despliegue y mantenimiento; y aplica el mismo criterio
  técnico a las demás metodologías. Cuando exista más de una, explica cómo se
  complementan dentro del proyecto. No menciones metodologías no seleccionadas y no
  incluyas instrucciones sobre corrección, generación del texto o uso de inteligencia
  artificial.
- Actividades corregidas: conserva exactamente la cantidad y el sentido de las
  actividades suministradas, pero nunca copies literalmente su redacción. Corrige
  ortografía, gramática, puntuación y concordancia; reformula cada actividad con
  lenguaje técnico, institucional y claro. Mejora su precisión cuando sea necesario
  para identificar correctamente la acción realizada, el proceso desarrollado o el
  componente intervenido, sin inventar tecnologías, resultados, cantidades,
  validaciones, estados, observaciones, entregables ni evidencias.
- Entregables corregidos: conserva exactamente la cantidad y el sentido de los
  entregables suministrados. Corrige ortografía, gramática y puntuación; mejora
  su redacción y precisión técnica para que cada elemento identifique claramente
  el producto, prototipo, sistema, componente, documento o desarrollo obtenido.
  No inventes funciones, tecnologías, cantidades, validaciones ni características
  que no estén presentes en la información original.
"""

    entrada_bloque_1 = contexto + """

ESTRUCTURA JSON OBLIGATORIA
{
  "introduccion": "300 a 340 palabras",
  "planteamiento_problema": "300 a 340 palabras",
  "objetivo_general": "una oración",
  "objetivos_especificos": ["objetivo 1", "objetivo 2", "objetivo 3", "objetivo 4"],
  "estado_arte_tecnica": "300 a 340 palabras con dos citas",
  "metodologia_desarrollo": "300 a 340 palabras",
  "actividades_corregidas": ["actividad corregida 1", "actividad corregida 2"],
  "entregables_corregidos": ["entregable mejorado 1", "entregable mejorado 2"]
}
"""

    bloque_1 = generar_json_openai(
        instrucciones=instrucciones_bloque_1,
        entrada=entrada_bloque_1,
        modelo=modelo_openai,
        temperature=0.15,
    )

    bloque_1_normalizado = normalizar_contenido(
        bloque_1,
        respaldo,
        datos,
    )
    actividades_corregidas = bloque_1_normalizado[
        "actividades_corregidas"
    ]
    entregables_corregidos = bloque_1_normalizado[
        "entregables_corregidos"
    ]

    contexto_bloque_2 = f"""
DATOS GENERALES
Talento: {datos.get('nombre_talento', '')}
Proyecto: {datos.get('nombre_proyecto', '')}
Código: {datos.get('codigo_proyecto', '')}
Experto: {datos.get('nombre_experto', '')}
Línea tecnológica: {datos.get('linea_tecnologica', '')}
TRL inicial: {datos.get('trl_inicial', '')}
TRL alcanzado: {datos.get('trl_alcanzado', '')}
TecnoParque: {datos.get('tecnoparque', '')}

DESCRIPCIÓN GENERAL DEL PROYECTO
{datos.get('descripcion_general_proyecto', '')}

INNOVACIÓN DEL PROYECTO
{datos.get('innovacion_proyecto_base', '')}

ACTIVIDADES CORREGIDAS
{actividades_en_texto(actividades_corregidas)}

ENTREGABLES CORREGIDOS Y MEJORADOS TÉCNICAMENTE
{actividades_en_texto(entregables_corregidos)}

IMPACTO DEL PROYECTO
{datos.get('impacto_proyecto_base', '')}

METODOLOGÍAS SELECCIONADAS
{metodologia}

FASES DE REFERENCIA
{fases_metodologias_en_texto(datos)}
"""

    instrucciones_bloque_2 = reglas_comunes + """
Genera los apartados 7, 9, 10, 11, 12 y 13.

REQUISITOS ESPECÍFICOS
- Desarrollo del proyecto: genera una introducción técnica breve y una entrada
  independiente por cada actividad corregida, conservando exactamente la cantidad y
  el orden de las actividades. Para cada actividad crea un título técnico, identifica
  la fase o fases metodológicas relacionadas y redacta una descripción técnica de
  130 a 220 palabras sobre su ejecución, propósito, procedimiento, decisiones,
  componentes, integraciones, revisiones o ajustes.
- Relaciona cada actividad con las fases reales de las metodologías seleccionadas.
  No asignes fases de metodologías que no fueron elegidas.
- No copies literalmente las actividades originales ni las corregidas. Utilízalas
  como insumo y reformula su contenido con lenguaje técnico, ortografía, gramática y
  precisión mejoradas, sin inventar tecnologías, resultados, pruebas, cantidades o
  validaciones.
- No agrupes varias actividades en una sola entrada y no omitas ninguna.
- No incluyas frases sobre el proceso de redacción, corrección, generación, prompt,
  inteligencia artificial o instrucciones recibidas. Presenta únicamente el contenido
  técnico final del proyecto.
- Resultados: utiliza exclusivamente los entregables corregidos como fuente
  conceptual y explica su relación con objetivos, actividades, innovación y TRL.
  No copies, enumeres ni reproduzcas literalmente los entregables.
- Viabilidad: analiza solo los aspectos técnicos, operativos, económicos,
  normativos, de adopción, sostenibilidad y escalabilidad pertinentes.
- Propiedad intelectual y transferencia: determina automáticamente los mecanismos
  realmente pertinentes en Colombia. No afirmes que exista un derecho concedido.
- Impacto: desarrolla el impacto suministrado y solo las dimensiones aplicables.
- Conclusiones: integra pertinencia, entregables, innovación, limitaciones y
  aprendizajes. Incluye exactamente la recomendación TRL suministrada.
"""

    entrada_bloque_2 = contexto_bloque_2 + f"""

RECOMENDACIÓN OBLIGATORIA DE CONTINUIDAD TRL
{recomendacion_continuidad_trl(datos.get('trl_alcanzado', ''))}

ESTRUCTURA JSON OBLIGATORIA
{{
  "desarrollo_proyecto": "introducción técnica de 120 a 190 palabras",
  "desarrollo_actividades": [
    {{
      "titulo": "nombre técnico corregido de la actividad",
      "fase_metodologica": "fase o fases y metodología relacionada",
      "descripcion_tecnica": "130 a 220 palabras sobre la ejecución técnica",
      "evidencia": "Insertar fotografías, capturas de pantalla, enlaces o soportes correspondientes a esta actividad."
    }}
  ],
  "resultados_obtenidos": "300 a 340 palabras",
  "analisis_viabilidad": "300 a 340 palabras",
  "propiedad_transferencia": "300 a 340 palabras",
  "impacto_proyecto": "300 a 340 palabras",
  "conclusiones": "300 a 340 palabras"
}}
"""

    bloque_2 = generar_json_openai(
        instrucciones=instrucciones_bloque_2,
        entrada=entrada_bloque_2,
        modelo=modelo_openai,
        temperature=0.15,
    )

    contenido: dict = {}
    if isinstance(bloque_1, dict):
        contenido.update(bloque_1)
    if isinstance(bloque_2, dict):
        contenido.update(bloque_2)

    contenido["actividades_corregidas"] = actividades_corregidas
    contenido["entregables_corregidos"] = entregables_corregidos

    normatividad = datos.get("normatividad_investigada", {})
    if not isinstance(normatividad, dict):
        normatividad = {}
    contenido["normatividad_aplicable"] = limpiar_texto(
        str(normatividad.get("normatividad_aplicable", ""))
    )
    contenido["normas_aplicables"] = normalizar_normas_aplicables(
        normatividad.get("normas_aplicables", [])
    )

    contenido["desarrollo_actividades"] = (
        normalizar_desarrollo_actividades(
            contenido.get("desarrollo_actividades", []),
            actividades_corregidas,
            datos,
        )
    )

    resultados_generados = limpiar_texto(
        str(contenido.get("resultados_obtenidos", ""))
    )

    if contiene_copia_textual(
        resultados_generados,
        entregables_originales,
    ):
        contenido["resultados_obtenidos"] = (
            reescribir_resultados_sin_copia_literal(
                texto_resultados=resultados_generados,
                entregables_corregidos=entregables_corregidos,
                entregables_originales=entregables_originales,
                datos=datos,
                modelo_openai=modelo_openai,
            )
        )

    contenido["referencias_bibliograficas"] = (
        referencias_bibliograficas_proyecto(datos)
    )
    contenido["anexos"] = anexos_manuales()

    contenido_normalizado = normalizar_contenido(
        contenido,
        respaldo,
        datos,
    )

    return contenido_normalizado


# =====================================================
# GENERACIÓN DEL DOCUMENTO OFICIAL
# =====================================================

def preparar_toc_para_word(documento: Document) -> None:
    """
    Deja la tabla de contenido configurada para que Microsoft Word recalcule
    automáticamente los números de página al abrir el archivo.

    No ejecuta LibreOffice, UNO, conversiones a PDF ni procesos externos.
    Esto evita los cierres nativos de Streamlit Cloud.
    """
    configurar_estilos_y_tabla_contenido(documento)
    marcar_actualizacion_campos(documento)



def generar_docx_informe_tecnico_final(datos: dict) -> str:
    plantilla = obtener_ruta_plantilla()
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    codigo_archivo = safe_filename(
        datos.get("codigo_proyecto", "proyecto")
    )
    nombre_archivo = f"Informe Final {codigo_archivo}.docx"
    ruta_salida = OUTPUT_DIR / nombre_archivo

    documento = Document(str(plantilla))

    eliminar_instrucciones_y_control_cambios(documento)
    renumerar_apartados_posteriores(documento)
    marcar_clasificacion(
        documento,
        datos.get("clasificacion_informacion", "Pública"),
    )
    llenar_tabla_informacion_general(documento, datos)

    fecha_portada = fecha_mes_anio_espanol(datos["fecha_entrega"])

    for parrafo in documento.paragraphs:
        if "junio 2026" in parrafo.text.casefold():
            escribir_parrafo(
                parrafo,
                fecha_portada,
                tamano=12,
                negrita=True,
                alineacion=WD_ALIGN_PARAGRAPH.CENTER,
            )
            break

    contenido = datos["contenido_informe"]

    mapa_reemplazos = [
        ("Introducción", contenido["introduccion"]),
        (
            "Planteamiento del problema",
            contenido["planteamiento_problema"],
        ),
        ("4.1 Objetivo General", contenido["objetivo_general"]),
        (
            "5. Estado del arte y estado de la técnica",
            contenido["estado_arte_tecnica"],
        ),
        (
            "6. Metodología de desarrollo",
            contenido["metodologia_desarrollo"],
        ),
        (
            "7. Desarrollo del proyecto",
            contenido["desarrollo_proyecto"],
        ),
        (
            "9. Resultados obtenidos",
            contenido["resultados_obtenidos"],
        ),
        (
            "10. Análisis de viabilidad",
            contenido["analisis_viabilidad"],
        ),
        (
            "11. Propiedad intelectual y transferencia tecnológica",
            contenido["propiedad_transferencia"],
        ),
        (
            "12. Impacto del proyecto",
            contenido["impacto_proyecto"],
        ),
        ("13. Conclusiones", contenido["conclusiones"]),
    ]

    destinos: dict[str, Paragraph] = {}

    for titulo, texto_apartado in mapa_reemplazos:
        encabezado = buscar_parrafo(documento, titulo)
        destino = parrafo_siguiente(encabezado)
        destino.style = "Normal"
        escribir_parrafo(
            destino,
            texto_apartado,
            tamano=11,
            alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        destinos[titulo] = destino

    ultimo_desarrollo = insertar_desarrollo_actividades(
        destinos["7. Desarrollo del proyecto"],
        contenido["desarrollo_actividades"],
    )

    insertar_normatividad(
        ultimo_desarrollo,
        contenido["normatividad_aplicable"],
        contenido["normas_aplicables"],
    )

    encabezado_objetivos_especificos = buscar_parrafo(
        documento,
        "4.2 Objetivos Específicos",
    )
    destino_objetivos = parrafo_siguiente(
        encabezado_objetivos_especificos
    )
    escribir_lista_en_parrafos(
        destino_objetivos,
        contenido["objetivos_especificos"],
    )

    # Elimina completamente el texto introductorio genérico de Objetivos.
    encabezado_objetivos = buscar_parrafo(
        documento,
        "Objetivos",
        coincidencia_exacta=True,
    )
    introduccion_objetivos = parrafo_siguiente(encabezado_objetivos)
    eliminar_parrafo(introduccion_objetivos)

    # Tabla institucional de resultados basada únicamente en los entregables.
    insertar_tabla_resultados(
        documento,
        destinos["9. Resultados obtenidos"],
        contenido["entregables_corregidos"],
    )

    encabezado_referencias = buscar_parrafo(
        documento,
        "14. Referencias bibliográficas",
    )
    destino_referencias = parrafo_siguiente(encabezado_referencias)
    escribir_lista_en_parrafos(
        destino_referencias,
        contenido["referencias_bibliograficas"],
    )

    encabezado_anexos = buscar_parrafo(
        documento,
        "15. Anexos",
        coincidencia_exacta=True,
    )
    # La plantilla trae un salto manual antes de Anexos. Se elimina para
    # evitar una página en blanco. Anexos continuará en la página disponible.
    parrafo_previo_xml = encabezado_anexos._p.getprevious()
    if parrafo_previo_xml is not None:
        for salto in parrafo_previo_xml.xpath(
            './/*[local-name()="br"]'
        ):
            padre = salto.getparent()
            if padre is not None:
                padre.remove(salto)

    encabezado_anexos.paragraph_format.page_break_before = False
    destino_anexos = parrafo_siguiente(encabezado_anexos)
    escribir_lista_en_parrafos(
        destino_anexos,
        contenido["anexos"],
    )

    documento.core_properties.title = (
        f"Informe Final {datos.get('codigo_proyecto', '')}"
    )
    documento.core_properties.subject = CODIGO_FORMATO_INFORME
    documento.core_properties.keywords = (
        "SENA, TecnoParque, Informe Final, GCDTP-F-023 V01"
    )

    preparar_toc_para_word(documento)

    # Limpieza final del salto heredado antes de Anexos, después de aplicar estilos.
    encabezado_anexos_final = buscar_parrafo(
        documento,
        "15. Anexos",
        coincidencia_exacta=True,
    )
    encabezado_anexos_final.paragraph_format.page_break_before = False
    previo_final = encabezado_anexos_final._p.getprevious()
    if previo_final is not None:
        for salto in previo_final.xpath('.//*[local-name()="br"]'):
            padre = salto.getparent()
            if padre is not None:
                padre.remove(salto)

    documento.save(str(ruta_salida))

    datos_json = serializar_datos_informe(datos)
    datos_json["ruta_docx"] = str(ruta_salida)

    guardar_datos_json(
        datos_json,
        nombre_archivo="datos_informe_tecnico_final.json",
    )

    return str(ruta_salida)


# =====================================================
# INTERFAZ STREAMLIT
# =====================================================

def render_informe_tecnico_final(
    modo_prueba: bool = True,
    modelo_openai: str = "gpt-4.1-mini",
) -> None:
    st.markdown("---")
    st.subheader("Informe Técnico Final")
    st.caption(
        f"{CODIGO_FORMATO_INFORME} · "
        f"{VERSION_INFORME_TECNICO_FINAL}"
    )

    st.info(
        "El formulario solicita los datos institucionales, la descripción general, "
        "los entregables, la innovación, la metodología utilizada, las actividades, "
        "la normatividad aplicable y el impacto. Las actividades se escriben en un solo campo "
        "y se articulan con la metodología para generar el desarrollo del proyecto. "
        "La tabla de contenido queda configurada para actualizarse automáticamente "
        "al abrir el archivo en Microsoft Word."
    )

    datos_key = "datos_informe_tecnico_final_formulario_minimo"
    ruta_key = "ruta_docx_informe_tecnico_final_formulario_minimo"

    if datos_key not in st.session_state:
        st.session_state[datos_key] = None

    if ruta_key not in st.session_state:
        st.session_state[ruta_key] = None

    with st.container():
        st.markdown("## 1. Información general del proyecto")

        col_1, col_2 = st.columns(2)

        with col_1:
            clasificacion_informacion = st.selectbox(
                "Clasificación de la información",
                options=CLASIFICACIONES_INFORMACION,
            )

            nombre_talento = st.text_input(
                "Talento que realiza el informe",
            )

            nombre_proyecto = st.text_area(
                "Nombre del Proyecto de Base Tecnológica",
                height=90,
            )

            codigo_proyecto = st.text_input(
                "Código de la idea",
                placeholder="Ejemplo: P2026-143440-18724",
            )

        with col_2:
            nombre_experto = st.text_input(
                "Experto del proyecto",
            )

            linea_tecnologica = st.text_input(
                "Línea tecnológica",
            )

            trl_inicial = st.selectbox(
                "TRL inicial",
                options=[f"TRL {numero}" for numero in range(1, 10)],
                index=4,
            )

            trl_alcanzado = st.selectbox(
                "TRL alcanzado",
                options=[f"TRL {numero}" for numero in range(1, 10)],
                index=5,
            )

            tecnoparque = st.text_input(
                "TecnoParque donde se desarrolló el proyecto",
                value="TecnoParque Nodo Angostura",
            )

            fecha_entrega = st.date_input(
                "Fecha de entrega",
                value=date.today(),
            )

        st.markdown("## Información técnica mínima")

        descripcion_general_proyecto = st.text_area(
            "Descripción general del proyecto",
            placeholder=(
                "Describe el origen, la necesidad, los usuarios, el propósito, "
                "los componentes, las tecnologías, el funcionamiento y el "
                "contexto de aplicación."
            ),
            height=340,
        )

        entregables_proyecto_base = st.text_area(
            "Entregables obtenidos",
            placeholder=(
                "Escribe un entregable por línea. Puedes registrar productos, "
                "prototipos, componentes, documentos, sistemas o desarrollos "
                "realmente obtenidos al cierre."
            ),
            height=210,
            help=(
                "El sistema corregirá la ortografía y mejorará la redacción y "
                "precisión técnica de cada entregable sin inventar características."
            ),
        )

        innovacion_proyecto_base = st.text_area(
            "Innovación del proyecto",
            placeholder=(
                "Explica el elemento diferencial, la mejora frente a alternativas "
                "existentes y el aporte técnico o funcional de la solución."
            ),
            height=190,
        )

        metodologias_seleccionadas = st.multiselect(
            "Metodología utilizada",
            options=METODOLOGIAS_DESARROLLO,
            help=(
                "Selecciona una o varias metodologías aplicadas durante el proyecto. "
                "Cuando elijas Otra, se habilitará un campo para escribirla."
            ),
            key="informe_final_metodologias_v1",
        )

        otra_metodologia = ""

        if "Otra" in metodologias_seleccionadas:
            otra_metodologia = st.text_input(
                "Escribe la metodología utilizada",
                placeholder=(
                    "Ejemplo: metodología propia de diseño y validación por etapas."
                ),
                key="informe_final_otra_metodologia_v1",
            )

        st.markdown("### Actividades desarrolladas")

        actividades_ejecutadas_texto = st.text_area(
            "Descripción de las actividades ejecutadas",
            placeholder=(
                "Escribe una actividad por línea. También puedes usar numeración, "
                "viñetas o separar las actividades con punto y coma.\n\n"
                "Ejemplo:\n"
                "1. Levantamiento de requerimientos con los usuarios.\n"
                "2. Diseño de la arquitectura de la solución.\n"
                "3. Desarrollo e integración de los componentes.\n"
                "4. Pruebas, ajustes y validación del prototipo."
            ),
            height=230,
            key="informe_final_actividades_texto_v1",
            help=(
                "El sistema corregirá la ortografía, mejorará la redacción y "
                "articulará estas actividades con la metodología seleccionada."
            ),
        )

        impacto_proyecto_base = st.text_area(
            "Impacto del proyecto",
            placeholder=(
                "Describe los beneficios, beneficiarios y efectos tecnológicos, "
                "sociales, económicos, ambientales o productivos identificados."
            ),
            height=220,
        )

        generar_contenido = st.button(
            (
                "Generar contenido en modo prueba"
                if modo_prueba
                else "Generar Informe Final con la API de OpenAI"
            )
        )

    if generar_contenido:
        campos_obligatorios = {
            "Talento que realiza el informe": nombre_talento,
            "Nombre del proyecto": nombre_proyecto,
            "Código de la idea": codigo_proyecto,
            "Experto del proyecto": nombre_experto,
            "Línea tecnológica": linea_tecnologica,
            "Descripción general del proyecto": descripcion_general_proyecto,
            "Entregables obtenidos": entregables_proyecto_base,
            "Innovación del proyecto": innovacion_proyecto_base,
            "Impacto del proyecto": impacto_proyecto_base,
        }

        if not validar_campos_obligatorios(campos_obligatorios):
            st.stop()

        metodologias_validas = [
            item
            for item in metodologias_seleccionadas
            if item != "Otra"
        ]

        if not metodologias_validas and not limpiar_texto(otra_metodologia):
            st.error(
                "Selecciona al menos una metodología o utiliza la opción Otra "
                "para escribirla manualmente."
            )
            st.stop()

        if (
            "Otra" in metodologias_seleccionadas
            and not limpiar_texto(otra_metodologia)
        ):
            st.error(
                "Escribe la metodología utilizada en el campo habilitado."
            )
            st.stop()

        entregables_validos = dividir_entregables(
            entregables_proyecto_base
        )

        if not entregables_validos:
            st.error(
                "Escribe al menos un entregable obtenido."
            )
            st.stop()

        actividades_validas = dividir_actividades(
            actividades_ejecutadas_texto
        )

        if not actividades_validas:
            st.error(
                "Escribe al menos una actividad desarrollada."
            )
            st.stop()

        datos_base = {
            "tipo_documento": "Informe Final",
            "codigo_formato": CODIGO_FORMATO_INFORME,
            "clasificacion_informacion": limpiar_texto(
                clasificacion_informacion
            ),
            "nombre_talento": limpiar_texto(nombre_talento),
            "nombre_proyecto": limpiar_texto(nombre_proyecto),
            "codigo_proyecto": limpiar_texto(codigo_proyecto),
            "nombre_experto": limpiar_texto(nombre_experto),
            "linea_tecnologica": limpiar_texto(linea_tecnologica),
            "trl_inicial": limpiar_texto(trl_inicial),
            "trl_alcanzado": limpiar_texto(trl_alcanzado),
            "tecnoparque": limpiar_texto(tecnoparque),
            "fecha_entrega": fecha_entrega,
            "fecha_entrega_texto": fecha_entrega.strftime("%d/%m/%Y"),
            "descripcion_general_proyecto": limpiar_texto(
                descripcion_general_proyecto
            ),
            "entregables_proyecto_base": limpiar_texto(
                entregables_proyecto_base
            ),
            "entregables_proyecto_lista": entregables_validos,
            "innovacion_proyecto_base": limpiar_texto(
                innovacion_proyecto_base
            ),
            "metodologias_seleccionadas": metodologias_seleccionadas,
            "otra_metodologia": limpiar_texto(otra_metodologia),
            "actividades_ejecutadas_base": actividades_validas,
            "impacto_proyecto_base": limpiar_texto(
                impacto_proyecto_base
            ),
            "modo_generacion": (
                "Prueba local"
                if modo_prueba
                else "API de OpenAI con búsqueda web"
            ),
            "version": VERSION_INFORME_TECNICO_FINAL,
        }
        datos_base["metodologia_inferida"] = (
            metodologias_en_texto(datos_base)
        )

        with st.spinner(
            "Investigando referentes reales, normatividad aplicable y generando "
            "los apartados del informe."
        ):
            try:
                if modo_prueba:
                    datos_base["referentes_estado_arte"] = (
                        referentes_modo_prueba(datos_base)
                    )
                    datos_base["normatividad_investigada"] = (
                        normatividad_modo_prueba(datos_base)
                    )
                    contenido = contenido_modo_prueba(datos_base)
                else:
                    datos_base["referentes_estado_arte"] = (
                        investigar_referentes_reales(
                            datos_base,
                            modelo_openai,
                        )
                    )
                    datos_base["normatividad_investigada"] = (
                        investigar_normatividad_aplicable(
                            datos_base,
                            modelo_openai,
                        )
                    )
                    contenido = generar_contenido_con_ia(
                        datos_base,
                        modelo_openai,
                    )
            except Exception as error:
                st.error(
                    f"No se pudo generar el contenido del informe: {error}"
                )
                st.stop()

        datos_base["contenido_informe"] = contenido
        st.session_state[datos_key] = datos_base
        st.session_state[ruta_key] = None

        st.success(
            "Contenido generado. Puedes revisarlo antes de crear el Word oficial."
        )

    datos = st.session_state.get(datos_key)

    if datos:
        contenido = datos["contenido_informe"]

        st.success(
            "Todos los apartados del informe fueron generados automáticamente. "
            "No es necesario diligenciar información adicional por sección."
        )

        st.markdown("### Resumen de generación")
        st.write(
            "**Metodología utilizada:**",
            metodologias_en_texto(datos),
        )
        st.write(
            "**Actividades procesadas para el desarrollo:**",
            len(contenido.get("actividades_corregidas", [])),
        )
        st.write(
            "**Entregables procesados para la tabla de resultados:**",
            len(contenido.get("entregables_corregidos", [])),
        )
        st.write(
            "**Normas y estándares identificados:**",
            len(contenido.get("normas_aplicables", [])),
        )
        st.caption(
            "El desarrollo del proyecto presenta cada actividad por separado, "
            "relacionada con las fases de las metodologías seleccionadas y con un "
            "espacio para incorporar evidencias en el documento Word."
        )

        st.markdown("### Vista previa de la tabla de resultados")
        tabla_resultados = [
            {
                "N.°": indice,
                "Descripción del entregable obtenido": entregable,
                "Evidencia": "Agregar enlace en Word",
            }
            for indice, entregable in enumerate(
                contenido["entregables_corregidos"],
                start=1,
            )
        ]
        st.dataframe(
            tabla_resultados,
            width="stretch",
            hide_index=True,
        )

        col_json, col_docx = st.columns(2)

        with col_json:
            datos_json_descarga = json.dumps(
                serializar_datos_informe(datos),
                ensure_ascii=False,
                indent=4,
            ).encode("utf-8")

            st.download_button(
                label="Descargar datos en JSON",
                data=datos_json_descarga,
                file_name="datos_informe_tecnico_final.json",
                mime="application/json",
                on_click="ignore",
                key="descargar_datos_informe_tecnico_final",
            )

        with col_docx:
            if st.button(
                "📄 Generar Word oficial GCDTP-F-023 V01",
                key="generar_docx_informe_tecnico_final_minimo",
            ):
                try:
                    ruta_docx = generar_docx_informe_tecnico_final(
                        datos
                    )
                    st.session_state[ruta_key] = ruta_docx
                    st.success(
                        "Documento Word generado correctamente."
                    )
                except Exception as error:
                    st.error(
                        f"No se pudo generar el documento Word: {error}"
                    )

        ruta_docx = st.session_state.get(ruta_key)

        if ruta_docx and Path(ruta_docx).exists():
            try:
                contenido_docx = Path(ruta_docx).read_bytes()
            except OSError as error:
                st.error(
                    f"No se pudo preparar el archivo para descarga: {error}"
                )
            else:
                st.download_button(
                    label="⬇️ Descargar Informe Final oficial",
                    data=contenido_docx,
                    file_name=Path(ruta_docx).name,
                    mime=(
                        "application/vnd.openxmlformats-"
                        "officedocument.wordprocessingml.document"
                    ),
                    on_click="ignore",
                    key="descargar_informe_final_oficial",
                )
